"""Generate sample-document.docx fixture cho Phase 7 MIGRATE-05 smoke E2E.

Run 1 lần:
    cd Hub_All/scripts/migrate/fixtures
    pip install python-docx
    python generate-sample.py
    # Output: sample-document.docx (5-10 KB)

Sau khi generate xong commit DOCX binary vào git (KHÔNG generate runtime — smoke
script reproducibility). KHÔNG dùng real production data (privacy + reproducibility).

Content phải có keyword `vaccin` + `dược` detectable cho smoke step 4-6 (search +
ask) assert non-empty result.
"""
from docx import Document
from pathlib import Path


def main() -> None:
    doc = Document()

    doc.add_heading("Sample Document — Smoke Test Phase 7", level=1)

    doc.add_paragraph(
        "Đây là tài liệu test ingest dùng cho smoke E2E Phase 7. Nội dung mô "
        "phỏng tài liệu y tế ngắn về vaccin và phòng dược cơ bản. KHÔNG dùng "
        "data production thật (privacy + reproducibility)."
    )

    doc.add_heading("Vaccin (vaccine)", level=2)
    doc.add_paragraph(
        "Vaccin là chế phẩm sinh học chứa kháng nguyên giảm độc lực hoặc bất "
        "hoạt, dùng để kích thích hệ miễn dịch sinh kháng thể chống lại tác "
        "nhân gây bệnh cụ thể. Tiêm chủng vaccin định kỳ giúp phòng ngừa các "
        "bệnh truyền nhiễm như cúm, sởi, bạch hầu, ho gà, uốn ván."
    )

    doc.add_heading("Phòng dược cơ bản", level=2)
    doc.add_paragraph(
        "Phòng dược trong cơ sở y tế chịu trách nhiệm cấp phát thuốc, kiểm "
        "kê tồn kho, theo dõi hạn sử dụng và phối hợp với bác sĩ điều trị "
        "để đảm bảo an toàn dược lý cho bệnh nhân. Dược sĩ chính cần có "
        "chứng chỉ hành nghề và cập nhật kiến thức theo quy định Bộ Y tế."
    )

    doc.add_heading("Tham khảo", level=2)
    doc.add_paragraph(
        "Tài liệu này chỉ phục vụ mục đích kiểm thử pipeline ingest + RAG "
        "search + ask + citation [N] cho Medinet Wiki v3.0 milestone "
        "closeout (Phase 7 MIGRATE-05). KHÔNG dùng làm tài liệu y khoa "
        "thực tế."
    )

    out_path = Path(__file__).parent / "sample-document.docx"
    doc.save(out_path)
    print(f"Generated: {out_path}")
    print(f"Size: {out_path.stat().st_size} bytes")


if __name__ == "__main__":
    main()
