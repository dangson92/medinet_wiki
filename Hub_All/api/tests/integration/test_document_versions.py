"""Integration test document_versions — Plan 05-04 v3.1 Phase 5 VER-05.

5 scenario E2E verify Plan 05-01..03 + Plan 02-01..04 RBAC carry forward:
- (1) create_version_via_reupload — seed doc + service snapshot → 2 row INSERT.
- (2) list_returns_ordered_desc — 3 mutation → GET DESC + envelope M2 shape.
- (3) restore_creates_new_version_append_only — POST /restore → v_max+1 + UPDATE doc.
- (4) hub_admin_cross_hub_versions_403 — hub_admin dmd POST /restore doc tdt → 403.
- (5) audit_forensic_chain — query 2 action 'document.version.{create,restore}' + payload.

Audit forensic chain (D-V3.1-Phase5-H LOCKED): scenario 1-3 emit audit; scenario 5
query asyncpg `payload->>'actor_role'` + `payload->>'document_id'` + `payload->>'version_number'`
+ `payload->>'restored_to'` exact match.

Reuse fixture từ conftest.py: postgres_container + redis_container + alembic_cfg +
app_with_auth + auth_client + admin_user|viewer_user + admin_token|viewer_token +
_login_get_token + GO_SEED_HASH (Phase 2 v3.0 + v3.1 ship).

Reuse pattern:
- test_smoke_e2e_v3_1_rbac.py (Plan 04-02 v3.1) — _seed_hub_admin_user + seed_hubs_dmd_tdt + audit poll.
- test_audit_actor_metadata.py (Plan 02-04 v3.1) — _wait_audit_row poll BackgroundTask timing.
- conftest.py existing fixture — KHÔNG redeclare.

Cocoindex re-extract: app.state.cocoindex_app default None ở test env (COCOINDEX_SKIP_SETUP=1
ở app_with_auth fixture) — router getattr(..., None) handle gracefully, KHÔNG cần monkeypatch.

DEVIATION (Rule 3) vs plan template: Plan template dùng `hub_app_factory("central")` (fake DSN,
KHÔNG apply migration, KHÔNG có Postgres thật connect) → scenario require DB schema sẽ FAIL.
Adapt sang `app_with_auth` fixture (test_smoke_e2e_v3_1_rbac.py:83-90 pattern) — auto-apply
alembic head + LifespanManager + truncate per-test isolation + share `auth_client` httpx
ASGITransport. Scenario 1 (service-layer) vẫn dùng `get_session()` public + ORM `Document`
fetch (KHÔNG động vào session factory private symbol — đã verify chỉ export get_session
+ init_engine + dispose_engine + get_engine ở app/db/session.py).

NOTE scenario 1 service-layer limitation: BE reupload endpoint chưa ship Phase 5
(per Plan 05-02 `<objective>` deferred note). Scenario 1 dùng public `get_session()`
async generator + fetch real ORM `Document` row qua `session.get(Document, doc_id)`
để exercise VER-02 snapshot service end-to-end.
"""
from __future__ import annotations

import asyncio
import json
import uuid
from pathlib import Path
from typing import Any

import httpx
import pytest
from docx import Document as DocxDocument  # python-docx — M2 dep (alias tránh clash ORM Document)
from sqlalchemy import text


# === Helper: seed hub row vào DB (id, slug, code, name, subdomain bắt buộc NOT NULL) ===

async def _seed_hub(
    *,
    code: str,
    name: str,
    slug: str | None = None,
    subdomain: str | None = None,
    hub_id: str | None = None,
) -> str:
    """INSERT 1 hub row + return hub_id UUID str.

    Cột bắt buộc NOT NULL theo migrations 0001 baseline + 0003 reconcile:
    - id (UUID PK gen_random_uuid)
    - slug (TEXT NOT NULL - Phase 2 baseline)
    - code (TEXT NOT NULL - Phase 5 reconcile; UNIQUE)
    - name (TEXT NOT NULL - Phase 2 baseline)
    - subdomain (TEXT NOT NULL - Phase 5 reconcile; server_default dropped sau backfill)
    - status (TEXT NOT NULL DEFAULT 'active')
    Default `slug=code` + `subdomain=code` mirror pattern Plan 04-02 v3.1
    seed_hubs_dmd_tdt fixture (test_smoke_e2e_v3_1_rbac.py).
    """
    from app.db.session import get_engine

    engine = get_engine()
    hub_id = hub_id or str(uuid.uuid4())
    slug = slug or code
    subdomain = subdomain or code
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO hubs (id, slug, code, name, subdomain, status, "
                "is_active, created_at, updated_at) "
                "VALUES (:id, :slug, :code, :name, :sub, 'active', TRUE, "
                "NOW(), NOW()) "
                "ON CONFLICT (code) DO NOTHING"
            ),
            {"id": hub_id, "slug": slug, "code": code, "name": name, "sub": subdomain},
        )
    return hub_id


# === Helper: seed document row vào DB ===

async def _seed_document(
    *,
    hub_id: str,
    filename: str,
    file_path: str,
    file_size: int = 1024,
    mime_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
) -> str:
    """INSERT 1 document row + return doc_id UUID str."""
    from app.db.session import get_engine

    engine = get_engine()
    doc_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO documents ("
                "id, hub_id, filename, file_path, file_size_bytes, mime_type, "
                "status, attempts, chunk_count, created_at, updated_at"
                ") VALUES ("
                ":id, :hub, :fn, :path, :size, :mime, 'completed', 0, 0, NOW(), NOW()"
                ")"
            ),
            {
                "id": doc_id, "hub": hub_id, "fn": filename, "path": file_path,
                "size": file_size, "mime": mime_type,
            },
        )
    return doc_id


# === Helper: seed hub admin user (mirror Plan 04-02 + test_smoke_e2e_v3_1_rbac pattern) ===

async def _seed_hub_admin_user(
    *,
    email: str,
    hub_id: str,
    password_hash: str,
) -> str:
    """Seed user role='editor' global + user_hubs.role='hub_admin' per-hub → return user_id.

    Pattern carry forward test_smoke_e2e_v3_1_rbac.py:41-74 (Plan 04-02 v3.1 ship).
    """
    from app.db.session import get_engine

    engine = get_engine()
    user_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO users (id, email, password_hash, full_name, role, "
                "phone, department, status, is_active, created_at, updated_at) "
                "VALUES (:id, :email, :hash, :name, 'editor', NULL, NULL, "
                "'active', TRUE, NOW(), NOW())"
            ),
            {
                "id": user_id, "email": email, "hash": password_hash,
                "name": email.split("@")[0],
            },
        )
        await conn.execute(
            text(
                "INSERT INTO user_hubs (user_id, hub_id, role, assigned_at) "
                "VALUES (:uid, :hid, 'hub_admin', NOW())"
            ),
            {"uid": user_id, "hid": hub_id},
        )
    return user_id


# === sample_docs fixture: 2 DOCX inline qua python-docx ===

@pytest.fixture
def sample_docs(tmp_path: Path) -> list[Path]:
    """Tạo 2 DOCX nhỏ (~5KB mỗi cái) inline qua python-docx — KHÔNG cần fixture file disk.

    Returns:
        [path_v1, path_v2] — 2 Path tới DOCX khác content (file_hash khác nhau).
    """
    doc1 = DocxDocument()
    doc1.add_paragraph("Sample document v1 — vaccin covid 2026 nội dung gốc")
    path1 = tmp_path / "sample-v1.docx"
    doc1.save(path1)

    doc2 = DocxDocument()
    doc2.add_paragraph("Sample document v2 — vaccin covid updated 2026-05 nội dung mới")
    path2 = tmp_path / "sample-v2.docx"
    doc2.save(path2)

    return [path1, path2]


# === Audit forensic helper poll-based (SQLAlchemy engine - tránh duplicate pool) ===

async def _wait_audit_row(
    *,
    action: str,
    document_id: str | None = None,
    timeout: float = 3.0,
) -> dict[str, Any] | None:
    """Poll audit_logs WHERE action (+ optional document_id) sau BackgroundTask emit.

    Carry forward pattern test_audit_actor_metadata.py:92-130 + test_smoke_e2e_v3_1_rbac.py:128-176
    + memory project_fastapi_bgtask_commit.

    Returns:
        payload dict | None nếu KHÔNG tìm thấy trong timeout.
    """
    from app.db.session import get_engine

    engine = get_engine()
    elapsed = 0.0
    while elapsed < timeout:
        async with engine.begin() as conn:
            if document_id is not None:
                result = await conn.execute(
                    text(
                        "SELECT payload FROM audit_logs "
                        "WHERE action = :action AND payload->>'document_id' = :doc_id "
                        "ORDER BY created_at DESC LIMIT 1"
                    ),
                    {"action": action, "doc_id": document_id},
                )
            else:
                result = await conn.execute(
                    text(
                        "SELECT payload FROM audit_logs WHERE action = :action "
                        "ORDER BY created_at DESC LIMIT 1"
                    ),
                    {"action": action},
                )
            row = result.fetchone()
        if row is not None:
            payload = row[0]
            if isinstance(payload, str):
                payload = json.loads(payload)
            return payload
        await asyncio.sleep(0.1)
        elapsed += 0.1
    return None


async def _assert_audit_version_metadata(
    *,
    action: str,
    document_id: str,
    expected_role: str,
    expected_hub_id: str | None,
    expected_version_number: int | None = None,
    expected_change_type: str | None = None,
    expected_restored_to: str | None = None,
    timeout: float = 3.0,
) -> dict[str, Any]:
    """Poll + assert audit row có metadata exact.

    Returns:
        payload dict (cho caller introspect thêm).

    Raises:
        AssertionError nếu KHÔNG tìm thấy row HOẶC metadata mismatch.
    """
    payload = await _wait_audit_row(
        action=action, document_id=document_id, timeout=timeout,
    )
    assert payload is not None, (
        f"Audit row not found within {timeout}s: action={action} document_id={document_id}"
    )

    assert payload.get("actor_role") == expected_role, (
        f"audit actor_role mismatch action={action}: expected={expected_role!r} "
        f"got={payload.get('actor_role')!r}"
    )
    assert payload.get("actor_hub_id") == expected_hub_id, (
        f"audit actor_hub_id mismatch action={action}: expected={expected_hub_id!r} "
        f"got={payload.get('actor_hub_id')!r}"
    )
    assert payload.get("document_id") == document_id, (
        f"audit document_id mismatch: expected={document_id} got={payload.get('document_id')}"
    )
    if expected_version_number is not None:
        assert payload.get("version_number") == expected_version_number, (
            f"audit version_number mismatch: expected={expected_version_number} "
            f"got={payload.get('version_number')}"
        )
    if expected_change_type is not None:
        assert payload.get("change_type") == expected_change_type, (
            f"audit change_type mismatch: expected={expected_change_type!r} "
            f"got={payload.get('change_type')!r}"
        )
    if expected_restored_to is not None:
        assert payload.get("restored_to") == expected_restored_to, (
            f"audit restored_to mismatch: expected={expected_restored_to} "
            f"got={payload.get('restored_to')}"
        )
    return payload


# === Scenario 1: create version via service snapshot direct (service-layer-only) ===

@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_scenario_1_create_version_via_reupload(
    app_with_auth: Any,
    sample_docs: list[Path],
) -> None:
    """Seed doc + 2 service snapshot call → assert 2 row + change_type + dedupe.

    NOTE service-layer limitation: BE reupload endpoint chưa ship Phase 5 (per Plan 05-02
    `<objective>` deferred note — reupload trigger wiring defer sau Phase 5). Scenario này
    test snapshot service trực tiếp qua public `get_session()` async generator + fetch real
    ORM `Document` row. KHÔNG động vào private session factory symbol (KHÔNG export public —
    sẽ ImportError). KHÔNG fabricate fake `_Doc` class (snapshot service yêu cầu real ORM
    attributes — fake class sẽ trigger AttributeError hoặc silent contract drift).
    """
    _ = app_with_auth  # trigger lifespan + migration + truncate

    # Seed hub + document trỏ tới sample_docs[0]
    hub_id = await _seed_hub(code="test1", name="Test Hub 1")
    doc_id = await _seed_document(
        hub_id=hub_id, filename="sample.docx", file_path=str(sample_docs[0]),
        file_size=sample_docs[0].stat().st_size,
    )

    # === v1 snapshot: lấy AsyncSession scope qua public get_session() async generator,
    # fetch real ORM Document row, call service.snapshot(), commit, exit scope ===
    from app.db.session import get_session
    from app.models.document import Document as DocumentORM
    from app.services import document_version_service

    async for session in get_session():
        doc_orm = await session.get(DocumentORM, uuid.UUID(doc_id))
        assert doc_orm is not None, f"Seeded document {doc_id} fetch fail"
        await document_version_service.snapshot(
            session=session, document=doc_orm, change_type="reupload",
            actor_role="admin", actor_hub_id=None,
        )
        # IMPORTANT: get_session() commit chỉ chạy SAU yield resume — nhưng
        # `break` triggers GeneratorExit raise → rollback path. Phải commit
        # tường minh TRƯỚC khi break để persist changes (memory
        # project_fastapi_bgtask_commit pattern carry forward).
        await session.commit()
        break

    # === v2 snapshot: UPDATE document.file_path → sample_docs[1], rồi snapshot lần 2 ===
    from app.db.session import get_engine
    engine = get_engine()
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "UPDATE documents SET file_path = :path, file_size_bytes = :size "
                "WHERE id = :id"
            ),
            {
                "path": str(sample_docs[1]),
                "size": sample_docs[1].stat().st_size,
                "id": doc_id,
            },
        )

    async for session in get_session():
        doc_orm = await session.get(DocumentORM, uuid.UUID(doc_id))
        assert doc_orm is not None
        assert doc_orm.file_path == str(sample_docs[1]), (
            "documents.file_path UPDATE phải reflect trong ORM fetch (session.get bypass cache)"
        )
        await document_version_service.snapshot(
            session=session, document=doc_orm, change_type="reupload",
            actor_role="admin", actor_hub_id=None,
        )
        # Commit tường minh trước break (xem comment v1 snapshot ở trên).
        await session.commit()
        break

    # Verify 2 row INSERT + version_number monotonic + dedupe by hash
    async with engine.begin() as conn:
        result = await conn.execute(
            text(
                "SELECT version_number, is_original, change_type, file_hash "
                "FROM document_versions WHERE document_id = :doc_id "
                "ORDER BY version_number"
            ),
            {"doc_id": doc_id},
        )
        rows = result.fetchall()

    assert len(rows) == 2, f"Expect 2 version row, got {len(rows)}"
    assert rows[0][0] == 1  # version_number
    assert rows[0][1] is True  # is_original
    assert rows[1][0] == 2
    assert rows[1][1] is False
    assert rows[1][2] == "reupload"  # change_type
    assert rows[0][3] != rows[1][3], (
        "v1 + v2 với file khác content → file_hash khác nhau (dedupe key D-V3.1-Phase5-A)"
    )


# === Scenario 2: list returns ordered DESC qua HTTP GET ===

@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_scenario_2_list_returns_ordered_desc(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    sample_docs: list[Path],
) -> None:
    """3 mutation → GET /api/documents/{id}/versions trả [v3, v2, v1] DESC + envelope M2."""
    headers = {"Authorization": f"Bearer {admin_token}"}

    hub_id = await _seed_hub(code="test2", name="Test Hub 2")
    doc_id = await _seed_document(
        hub_id=hub_id, filename="sample.docx", file_path=str(sample_docs[0]),
    )

    # Seed 3 versions trực tiếp DB (test isolation — tránh service complexity)
    from app.db.session import get_engine
    engine = get_engine()
    async with engine.begin() as conn:
        for v_num, change_type in [(1, "reupload"), (2, "reextract"), (3, "content_edit")]:
            await conn.execute(
                text(
                    "INSERT INTO document_versions "
                    "(document_id, version_number, is_original, name, file_type, file_size, "
                    " file_path, change_type) "
                    "VALUES (:doc, :v, :is_orig, 'sample.docx', 'docx', 1024, :path, :ct)"
                ),
                {
                    "doc": doc_id, "v": v_num, "is_orig": (v_num == 1),
                    "path": str(sample_docs[0]), "ct": change_type,
                },
            )

    # GET /versions → envelope + DESC order
    res = await auth_client.get(f"/api/documents/{doc_id}/versions", headers=headers)
    assert res.status_code == 200, f"GET fail: {res.status_code} {res.text}"
    data = res.json()
    assert data["success"] is True
    assert data["error"] is None
    assert data["meta"] is None
    versions = data["data"]["versions"]
    assert len(versions) == 3
    assert [v["version_number"] for v in versions] == [3, 2, 1], (
        f"DESC order expected [3, 2, 1], got {[v['version_number'] for v in versions]}"
    )
    assert versions[2]["is_original"] is True  # v1 = original
    assert versions[0]["change_type"] == "content_edit"  # v3 latest


# === Scenario 3: restore creates new version append-only ===

@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_scenario_3_restore_creates_new_version_append_only(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    sample_docs: list[Path],
) -> None:
    """POST /restore từ v1 → v_max+1 INSERT change_type='restore' + total tăng 1 (D-V3.1-Phase5-D)."""
    headers = {"Authorization": f"Bearer {admin_token}"}

    hub_id = await _seed_hub(code="test3", name="Test Hub 3")
    # documents.file_path trỏ tới sample_docs[1] (v2 file)
    doc_id = await _seed_document(
        hub_id=hub_id, filename="sample.docx", file_path=str(sample_docs[1]),
        file_size=sample_docs[1].stat().st_size,
    )

    # Seed v1 với file_path = sample_docs[0]
    from app.db.session import get_engine
    engine = get_engine()
    v1_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO document_versions "
                "(id, document_id, version_number, is_original, name, file_type, file_size, "
                " file_path, change_type) "
                "VALUES (:id, :doc, 1, true, 'sample-v1.docx', 'docx', :size, :path, 'reupload')"
            ),
            {
                "id": v1_id, "doc": doc_id,
                "size": sample_docs[0].stat().st_size,
                "path": str(sample_docs[0]),
            },
        )

    # POST /restore v1 — assert HTTP 200 + envelope success
    res = await auth_client.post(
        f"/api/documents/{doc_id}/versions/{v1_id}/restore", headers=headers,
    )
    assert res.status_code == 200, f"POST /restore fail: {res.status_code} {res.text}"
    data = res.json()
    assert data["success"] is True

    # Verify total versions = 2 (v1 + v2_restore), KHÔNG xoá v1 (append-only)
    async with engine.begin() as conn:
        result = await conn.execute(
            text(
                "SELECT version_number, change_type, file_path FROM document_versions "
                "WHERE document_id = :doc ORDER BY version_number"
            ),
            {"doc": doc_id},
        )
        rows = result.fetchall()

    assert len(rows) == 2, f"Append-only: expect 2 row, got {len(rows)}"
    assert rows[0][0] == 1
    assert rows[1][0] == 2
    assert rows[1][1] == "restore", (
        f"v2 phải có change_type='restore' (D-V3.1-Phase5-D), got {rows[1][1]!r}"
    )

    # === BLOCKER 2 fix: append-only D-V3.1-Phase5-D LOCKED contract verify ===
    # Plan 05-02 contract: snapshot TRƯỚC khi UPDATE documents (immutable history).
    # Restore-marker row (v2) PHẢI capture PRE-restore documents.file_path
    # = sample_docs[1] (file đang attach trước khi user khôi phục về v1).
    # Nếu executor implement restore khác (capture target state thay vì
    # pre-restore current state) thì assertion này sẽ catch — defense
    # against silent contract drift trong Plan 05-02 implementation.
    assert rows[1][2] == str(sample_docs[1]), (
        "Restore-marker row (v2) phải capture PRE-restore documents.file_path "
        "= sample_docs[1] (D-V3.1-Phase5-D LOCKED append-only — snapshot trước UPDATE), "
        f"got {rows[1][2]!r}"
    )
    assert rows[0][2] == str(sample_docs[0]), (
        "Original v1 row file_path KHÔNG bị thay đổi sau restore "
        "(append-only history immutable), "
        f"got {rows[0][2]!r}"
    )

    # Verify documents.file_path UPDATE = v1.file_path
    async with engine.begin() as conn:
        result = await conn.execute(
            text("SELECT file_path FROM documents WHERE id = :id"),
            {"id": doc_id},
        )
        doc_row = result.fetchone()
    assert doc_row[0] == str(sample_docs[0]), (
        f"documents.file_path phải UPDATE về v1 path, got {doc_row[0]!r}"
    )


# === Scenario 4: hub_admin cross-hub restore → 403 HUB_ADMIN_REQUIRED ===

@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_scenario_4_hub_admin_cross_hub_versions_403(
    auth_client: httpx.AsyncClient,
    app_with_auth: Any,
    sample_docs: list[Path],
) -> None:
    """Hub_admin dmd POST /restore của doc thuộc hub tdt → 403 HUB_ADMIN_REQUIRED envelope.

    R-V3.1-2 mitigation chain Phase 5 — BE Layer 3 authoritative defense in depth.
    """
    _ = app_with_auth  # trigger lifespan + migration + truncate
    from tests.integration.conftest import GO_SEED_HASH, GO_SEED_HASH_PLAINTEXT

    # Seed CẢ 2 hub thật (dmd + tdt) — scenario test cross-hub 403 defense in depth.
    hub_dmd_id = await _seed_hub(code="dmd", name="Đỗ Minh Đường")
    hub_tdt_id = await _seed_hub(code="tdt", name="Thuốc Dân Tộc")

    # Seed hub_admin dmd
    hub_admin_email = f"hadmin-dmd-{uuid.uuid4().hex[:8]}@medinet.vn"
    await _seed_hub_admin_user(
        email=hub_admin_email, hub_id=hub_dmd_id, password_hash=GO_SEED_HASH,
    )

    # Seed doc tdt + 1 version
    doc_tdt_id = await _seed_document(
        hub_id=hub_tdt_id, filename="tdt-doc.docx", file_path=str(sample_docs[0]),
        file_size=sample_docs[0].stat().st_size,
    )
    from app.db.session import get_engine
    engine = get_engine()
    v1_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO document_versions "
                "(id, document_id, version_number, is_original, name, file_type, file_size, "
                " file_path, change_type) "
                "VALUES (:id, :doc, 1, true, 'tdt-doc.docx', 'docx', 1024, :path, 'reupload')"
            ),
            {"id": v1_id, "doc": doc_tdt_id, "path": str(sample_docs[0])},
        )

    # Login hub_admin dmd
    login_res = await auth_client.post(
        "/api/auth/login",
        json={"email": hub_admin_email, "password": GO_SEED_HASH_PLAINTEXT},
    )
    assert login_res.status_code == 200, f"Login hub_admin fail: {login_res.text}"
    token = login_res.json()["data"]["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # POST /restore doc tdt → 403 HUB_ADMIN_REQUIRED (cross-hub reject defense in depth)
    res = await auth_client.post(
        f"/api/documents/{doc_tdt_id}/versions/{v1_id}/restore", headers=headers,
    )
    assert res.status_code == 403, (
        f"Cross-hub POST /restore should be 403, got {res.status_code} {res.text}"
    )
    envelope = res.json()
    assert envelope["success"] is False
    assert envelope["error"]["code"] == "HUB_ADMIN_REQUIRED", (
        f"Cross-hub reject envelope code phải HUB_ADMIN_REQUIRED, got {envelope['error']}"
    )


# === Scenario 5: audit forensic chain verified ===

@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_scenario_5_audit_forensic_chain(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    sample_docs: list[Path],
) -> None:
    """Query audit_logs.payload->>'actor_role' + 'document_id' + 'version_number' + 'restored_to'.

    D-V3.1-Phase5-H LOCKED: 2 action 'document.version.create' + 'document.version.restore'
    với payload nest đầy đủ metadata cho forensic.
    """
    headers = {"Authorization": f"Bearer {admin_token}"}

    hub_id = await _seed_hub(code="test5", name="Test Hub 5")
    doc_id = await _seed_document(
        hub_id=hub_id, filename="audit-doc.docx", file_path=str(sample_docs[1]),
        file_size=sample_docs[1].stat().st_size,
    )

    # Seed v1 manual (file_path = sample_docs[0])
    from app.db.session import get_engine
    engine = get_engine()
    v1_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO document_versions "
                "(id, document_id, version_number, is_original, name, file_type, file_size, "
                " file_path, change_type) "
                "VALUES (:id, :doc, 1, true, 'audit-v1.docx', 'docx', :size, :path, 'reupload')"
            ),
            {
                "id": v1_id, "doc": doc_id,
                "size": sample_docs[0].stat().st_size,
                "path": str(sample_docs[0]),
            },
        )

    # POST /restore v1 → emit 2 action ('create' restore version + 'restore')
    res = await auth_client.post(
        f"/api/documents/{doc_id}/versions/{v1_id}/restore", headers=headers,
    )
    assert res.status_code == 200, f"POST /restore fail: {res.text}"

    # Wait audit forensic + verify
    create_payload = await _assert_audit_version_metadata(
        action="document.version.create",
        document_id=doc_id,
        expected_role="admin",
        expected_hub_id=None,
        expected_change_type="restore",  # snapshot trước UPDATE với change_type='restore'
        timeout=3.0,
    )
    assert create_payload["version_number"] == 2, (
        f"Create version sau restore phải v_max+1 = 2, got {create_payload['version_number']}"
    )

    restore_payload = await _assert_audit_version_metadata(
        action="document.version.restore",
        document_id=doc_id,
        expected_role="admin",
        expected_hub_id=None,
        expected_restored_to=v1_id,
        timeout=3.0,
    )
    assert restore_payload["version_number"] == 1, (
        f"Restore action payload.version_number phải bằng v_target=1, "
        f"got {restore_payload['version_number']}"
    )

    # Verify 2 distinct action có trong DB
    async with engine.begin() as conn:
        result = await conn.execute(
            text(
                "SELECT DISTINCT action FROM audit_logs "
                "WHERE payload->>'document_id' = :doc_id "
                "ORDER BY action"
            ),
            {"doc_id": doc_id},
        )
        action_names = sorted([row[0] for row in result.fetchall()])

    assert "document.version.create" in action_names, (
        f"Missing 'document.version.create' in audit, got {action_names}"
    )
    assert "document.version.restore" in action_names, (
        f"Missing 'document.version.restore' in audit, got {action_names}"
    )
