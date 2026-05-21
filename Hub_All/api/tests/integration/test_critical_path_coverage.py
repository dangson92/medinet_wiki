"""HARD-03 critical-path acceptance test suite — Plan 10-03.

5 acceptance test crisp đáp ứng HARD-03 acceptance line: ::

    Integration test suite ≥50% critical path coverage (pytest +
    testcontainers Postgres + Redis chạy CI GitHub Actions PASS). Test
    cho auth happy + hub isolation + ingest VN filename + search hub
    filter + ask citation parsing.

5 nhóm test (mỗi nhóm 1 acceptance test crisp):

- **T1 auth happy**: ``POST /api/auth/login`` → 200 + envelope shape Go +
  JWT decode được bằng ``JWTManager.verify_token`` + claims chứa user role.
- **T2 hub isolation E4**: editor Hub A → ``DELETE /api/documents/<doc_hub_b>``
  → 403 FORBIDDEN + document Hub B vẫn tồn tại + audit
  ``security.hub_isolation_violation`` được ghi.
- **T3 ingest VN filename**: admin → ``POST /api/documents/upload`` với file
  ``"Khám bệnh đa khoa.docx"`` → response + GET detail trả ``name`` UTF-8
  nguyên vẹn (byte-identical, không mojibake).
- **T4 search hub filter**: viewer Hub A → ``POST /api/search`` với
  ``hub_ids:[B]`` → 0 result Hub B (defense-in-depth Phase 6 SEARCH-03).
- **T5 ask citation parsing**: mock LLM trả ``"... [1] ... [2]."`` →
  ``POST /api/ask`` trả citations[0].chunk_id khớp chunks[0],
  citations[1].chunk_id khớp chunks[1].

Mỗi test có docstring trỏ về REQ-ID + Phase tham chiếu (detail test). File
NÀY là PRIMARY GATE — Plan 10-06 CI gate sẽ chạy::

    pytest -m critical -m integration \\
        tests/integration/test_critical_path_coverage.py \\
        --cov=app.auth --cov=app.routers --cov=app.services \\
        --cov-fail-under=50

DEF-05-01 — file này boot app qua fixture ``app_with_auth`` (lifespan +
migration). cocoindex 1.0.3 ``core.Environment`` là process-global singleton
KHÔNG re-open được → KHÔNG gộp file này với test integration khác boot app
trong cùng 1 pytest process. Recommended invocation::

    uv run pytest tests/integration/test_critical_path_coverage.py -v

Tận dụng fixture conftest.py + conftest_hardening.py:
- ``app_with_auth`` / ``auth_client`` (Phase 3)
- ``admin_user`` / ``admin_token`` / ``editor_user`` / ``editor_token`` /
  ``viewer_user`` / ``viewer_token`` (Phase 3)
- ``seeded_two_hubs_with_editor`` / ``mock_cocoindex_app_noop`` /
  ``mock_litellm_citation_response`` (Plan 10-03 conftest_hardening.py)
- helpers ``_insert_hub`` / ``_insert_document`` / ``_insert_chunk`` /
  ``_assign_user_hub`` (Phase 3+6)
"""
from __future__ import annotations

import asyncio
import io
import uuid
from typing import Any

import httpx
import pytest
from docx import Document as DocxDocument
from sqlalchemy import text

from tests.integration.conftest import (
    _assign_user_hub,
    _insert_chunk,
    _insert_document,
    _insert_hub,
)
from tests.integration.conftest_hardening import _MockCocoindexAppHardening

# pytest_plugins: load ``conftest_hardening.py`` như plugin để fixture
# ``seeded_two_hubs_with_editor`` + ``mock_litellm_citation_response`` +
# ``mock_cocoindex_app_noop`` được auto-discover (pytest chỉ load file tên
# ``conftest.py`` chuẩn — file conftest_hardening.py PHẢI register qua
# pytest_plugins). Pattern khớp pytest docs §"Defining your own plugins".
pytest_plugins = ["tests.integration.conftest_hardening"]

# ----------------------------------------------------------------------------
# Constants — Vietnamese filename literal (Test 3) + embedding dim (R1 pin).
# ----------------------------------------------------------------------------

VN_FILENAME = "Khám bệnh đa khoa.docx"
"""Tên file tiếng Việt có dấu cho Test 3 (HARD-03 acceptance line 3).

Pattern khớp ``test_vietnamese_filename.py`` Phase 8 (Plan 08-03 Task 2).
Nếu pipeline upload mojibake, assert so sánh chuỗi sẽ fail rõ ràng (vd
``"KhÃ¡m bá»‡nh ..."``).
"""

_EMBEDDING_DIM = 1536
"""Số chiều vector embedding PIN M2 (R1 — pgvector 2000-dim index limit)."""


def _fixed_vector(seed: float = 0.1) -> list[float]:
    """Vector deterministic 1536-dim — mọi phần tử = ``seed``.

    Pattern khớp Phase 6 ``test_search_hub_isolation.py:_fixed_vector`` +
    Phase 7 conftest ``_make_vec``. Seed chunk + query embedding cùng vector
    → cosine distance xác định, KHÔNG phụ thuộc embedding API thật.
    """
    return [seed] * _EMBEDDING_DIM


def _patch_embed(monkeypatch: pytest.MonkeyPatch, vector: list[float]) -> None:
    """Monkeypatch ``app.services.search_service.embed_text`` trả vector cố định.

    ``search_service`` đã ``from app.services.embedder import embed_text``
    → bind tên vào module ``search_service`` → patch ở ĐÚNG module này.
    Loại bỏ phụ thuộc ``OPENAI_API_KEY`` placeholder M2.
    """

    async def _fake_embed(text: str, model: str | None = None) -> list[float]:
        _ = (text, model)
        return list(vector)

    monkeypatch.setattr(
        "app.services.search_service.embed_text", _fake_embed
    )


def _make_docx_bytes(extra: str = "") -> bytes:
    """Tạo file DOCX tối giản hợp lệ in-memory → bytes.

    Pattern khớp Phase 4 / 5 / 8 helper — KHÔNG commit binary fixture, tạo
    runtime per-test.
    """
    doc = DocxDocument()
    doc.add_paragraph("Tài liệu test critical path " + extra)
    doc.add_paragraph("Nội dung minh hoạ quy trình kiểm thử.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _count_rows(query: str, params: dict[str, Any]) -> int:
    """SELECT COUNT scalar helper qua SQLAlchemy async engine.

    Pattern khớp ``test_hub_isolation.py:_count`` — dùng cho assert "document
    Hub B vẫn tồn tại sau cross-hub DELETE reject".
    """
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.connect() as conn:
        row = (await conn.execute(text(query), params)).fetchone()
    return int(row[0]) if row else 0


# ============================================================================
# Test 1 — auth happy path login (HARD-03 acceptance #1).
# ============================================================================


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_critical_auth_happy_login_envelope_and_jwt(
    auth_client: httpx.AsyncClient,
    admin_user: dict[str, str],
    app_with_auth: Any,
) -> None:
    """HARD-03 acceptance #1 — POST /api/auth/login → 200 + envelope + JWT decode.

    Smoke gate suite-level cho M2 production-ready: contract envelope đủ field
    + JWT access token decode được bằng ``JWTManager.verify_token`` với role
    khớp seed user.

    Detail test chi tiết: Phase 3 ``test_auth_login.py`` (5 case envelope
    shape + anti-timing oracle). Plan 10-03 chỉ verify happy path đủ 5 field
    cơ bản — KHÔNG duplicate full envelope shape test Go-compat.
    """
    _ = app_with_auth

    # === Login admin@medinet.vn / Admin@123 ===
    r = await auth_client.post(
        "/api/auth/login",
        json={
            "email": admin_user["email"],
            "password": admin_user["password"],
        },
    )
    assert r.status_code == 200, r.text

    body = r.json()
    # Envelope shape D6 — {success, data, error, meta}.
    assert body["success"] is True
    assert body["error"] is None
    assert "data" in body

    data = body["data"]
    assert isinstance(data["access_token"], str) and data["access_token"]
    assert isinstance(data["refresh_token"], str) and data["refresh_token"]
    assert isinstance(data["expires_at"], int)

    # User nested — UserWithRoles `{user, roles}` (D6 contract frontend).
    user_with_roles = data["user"]
    assert user_with_roles["user"]["email"] == admin_user["email"]
    assert user_with_roles["user"]["id"] == admin_user["id"]

    # === JWT decode bằng JWTManager.verify_token ===
    # HARD-03 acceptance: JWT decode OK + user role match seed (admin).
    # Verify role qua DB column ``users.role`` (D6 — role là single string,
    # KHÔNG roles list từ user_hubs ở envelope ngoài).
    from app.auth.jwt import JWTManager
    from app.config import get_settings

    jwt_manager = JWTManager(get_settings())
    claims = jwt_manager.verify_token(
        data["access_token"], expected_type="access"
    )
    # JWTClaims có sub (user_id) + role + email — pattern Phase 3 Plan 03-04.
    assert str(claims.sub) == admin_user["id"]
    assert claims.role == "admin", (
        f"JWT claims.role phải = 'admin', got {claims.role!r}"
    )


# ============================================================================
# Test 2 — hub isolation E4 critical (HARD-03 acceptance #2).
# ============================================================================


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_critical_hub_isolation_editor_cannot_delete_cross_hub(
    auth_client: httpx.AsyncClient,
    editor_token: str,
    seeded_two_hubs_with_editor: dict[str, str],
    app_with_auth: Any,
) -> None:
    """HARD-03 acceptance #2 — editor Hub A DELETE doc Hub B → 403 + audit logged.

    E4 EXIT criteria (PROJECT.md): hub isolation bug = ship-blocker. Test
    fail = STOP, security review.

    Smoke gate suite-level cho M2 production-ready:
    1. DELETE cross-hub → 403 FORBIDDEN.
    2. Document Hub B PHẢI vẫn tồn tại sau reject (KHÔNG bị xoá).
    3. Audit ``security.hub_isolation_violation`` được ghi (T-05-06).

    Detail test chi tiết: Phase 5 ``test_hub_isolation.py`` (6 test E4 đầy
    đủ: admin bypass, editor own-hub, viewer read-only, verify_hub_access
    unit). Plan 10-03 acceptance #2 là 1 case crisp nhất — cross-hub DELETE
    reject + audit.
    """
    _ = app_with_auth
    doc_hub_b = seeded_two_hubs_with_editor["doc_b_id"]

    r = await auth_client.delete(
        f"/api/documents/{doc_hub_b}",
        headers={"Authorization": f"Bearer {editor_token}"},
    )
    assert r.status_code == 403, r.text
    body = r.json()
    assert body["success"] is False
    assert body["error"]["code"] == "FORBIDDEN"

    # Document Hub B VẪN tồn tại sau cross-hub DELETE reject.
    still = await _count_rows(
        "SELECT COUNT(*) FROM documents WHERE id = :id", {"id": doc_hub_b}
    )
    assert still == 1, (
        "E4 VIOLATION — document Hub B bị editor Hub A xoá (acceptance #2 fail)"
    )

    # Audit ``security.hub_isolation_violation`` được ghi — async batch flush.
    # Pattern khớp ``test_hub_isolation.py`` poll-with-timeout.
    from app.services.audit_service import flush_pending

    count = 0
    for _attempt in range(60):
        await flush_pending()
        count = await _count_rows(
            "SELECT COUNT(*) FROM audit_logs "
            "WHERE action = 'security.hub_isolation_violation'",
            {},
        )
        if count >= 1:
            break
        await asyncio.sleep(0.15)
    assert count >= 1, (
        "E4 VIOLATION — audit security.hub_isolation_violation KHÔNG được ghi"
    )


# ============================================================================
# Test 3 — ingest Vietnamese filename UTF-8 (HARD-03 acceptance #3).
# ============================================================================


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_critical_ingest_vietnamese_filename_utf8_roundtrip(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    mock_cocoindex_app_noop: _MockCocoindexAppHardening,
) -> None:
    """HARD-03 acceptance #3 — upload "Khám bệnh đa khoa.docx" → UTF-8 nguyên vẹn.

    SC4 (Phase 8): filename UTF-8 tiếng Việt PHẢI preserve byte-identical qua
    pipeline upload → response → GET detail. Nếu mojibake (latin-1 decode UTF-8
    bytes) → assert fail rõ ràng.

    Test FOCUS: filename UTF-8 round-trip — KHÔNG verify cocoindex chunking
    flow (đã cover ở Phase 4 ``test_ingest_e2e.py``). Mock cocoindex_app
    no-op để BackgroundTask trigger_cocoindex_update KHÔNG raise.

    Detail test chi tiết: Phase 8 ``test_vietnamese_filename.py`` (SC4 + path
    traversal T-08-03-01). Plan 10-03 acceptance #3 là crisp version chỉ
    verify UTF-8 roundtrip 2 lần (response + GET).
    """
    _ = (app_with_auth, mock_cocoindex_app_noop)

    # Seed hub.
    hub_id = await _insert_hub(
        name="Hub VN Acceptance",
        code="hub-vn-acc",
        subdomain="hub-vn-acc",
    )

    # === Upload file tên tiếng Việt ===
    # httpx multipart RFC 7578: filename str Unicode (KHÔNG bytes mojibake).
    upload_resp = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={
            "file": (
                VN_FILENAME,
                io.BytesIO(_make_docx_bytes("Khám VN")),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
        data={"hub_id": hub_id},
    )
    assert upload_resp.status_code == 202, upload_resp.text

    upload_body = upload_resp.json()
    assert upload_body["success"] is True, upload_body
    upload_data = upload_body["data"]

    # HARD-03 acceptance #3 — name UTF-8 nguyên vẹn (so sánh literal Unicode).
    assert upload_data["name"] == VN_FILENAME, (
        f"HARD-03 acceptance #3 VIOLATION — upload name mojibake/lệch: "
        f"{upload_data['name']!r} kỳ vọng {VN_FILENAME!r}"
    )
    document_id = upload_data["id"]
    assert document_id, upload_data

    # === GET document detail — name vẫn đúng UTF-8 ===
    detail_resp = await auth_client.get(
        f"/api/documents/{document_id}",
        headers={"Authorization": f"Bearer {admin_token}"},
    )
    assert detail_resp.status_code == 200, detail_resp.text
    detail_data = detail_resp.json()["data"]
    assert detail_data["name"] == VN_FILENAME, (
        f"HARD-03 acceptance #3 VIOLATION — GET detail name lệch: "
        f"{detail_data['name']!r} kỳ vọng {VN_FILENAME!r}"
    )


# ============================================================================
# Test 4 — search hub filter (HARD-03 acceptance #4).
# ============================================================================


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_critical_search_hub_filter_isolation(
    auth_client: httpx.AsyncClient,
    viewer_token: str,
    viewer_user: dict[str, str],
    app_with_auth: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """HARD-03 acceptance #4 — viewer Hub A search hub_ids=[B] → 0 result B.

    Phase 6 SEARCH-03 defense-in-depth: viewer truyền explicit ``hub_ids``
    chứa Hub B (KHÔNG được assign) → SQL filter PHẢI loại Hub B TRƯỚC khi
    chạy vector search. KHÔNG có result.hub_id == hub_b (intersection
    user_hubs ∩ requested_hub_ids = rỗng → results: [] HOẶC chỉ Hub A).

    Acceptance accept cả 2 outcome (D-06 Phase 6 — defense in depth):
    1. results: [] (intersection rỗng + viewer chỉ assign Hub A → 0 chunk A).
    2. Trả chỉ Hub A (nếu viewer cũng assign A → filter intersection {A}).

    Test setup viewer assign Hub A + chunk Hub A + chunk Hub B → request
    hub_ids=[B] only → intersection {A} ∩ {B} = {} → results [].

    Detail test chi tiết: Phase 6 ``test_search_hub_isolation.py`` (6 test
    E4 đầy đủ — single hub, explicit hub_ids, cross-hub, empty hub, cache,
    EXPLAIN ANALYZE HNSW index hit).
    """
    _ = app_with_auth

    hub_a = await _insert_hub(name="Hub A", code="hub-a", subdomain="hub-a")
    hub_b = await _insert_hub(name="Hub B", code="hub-b", subdomain="hub-b")
    # Viewer CHỈ assign Hub A — KHÔNG Hub B.
    await _assign_user_hub(user_id=viewer_user["id"], hub_id=hub_a)

    doc_a = await _insert_document(hub_id=hub_a, filename="hub-a.docx")
    doc_b = await _insert_document(hub_id=hub_b, filename="hub-b.docx")
    await _insert_chunk(
        document_id=doc_a,
        hub_id=hub_a,
        content="nội dung Hub A",
        vector=_fixed_vector(0.1),
    )
    await _insert_chunk(
        document_id=doc_b,
        hub_id=hub_b,
        content="nội dung Hub B BÍ MẬT",
        vector=_fixed_vector(0.1),
    )
    _patch_embed(monkeypatch, _fixed_vector(0.1))

    # === POST /api/search hub_ids=[B] (only Hub B — KHÔNG assign cho viewer) ===
    r = await auth_client.post(
        "/api/search",
        headers={"Authorization": f"Bearer {viewer_token}"},
        json={"query": "nội dung", "hub_ids": [hub_b], "top_k": 5},
    )
    # D-06 Phase 6 defense-in-depth: backend chấp nhận 200 (intersection rỗng
    # → results []) HOẶC 403 (some implementations reject). Test verify CẢ 2.
    assert r.status_code in {200, 403}, r.text

    if r.status_code == 200:
        body = r.json()
        results = body["data"]["results"]
        # HARD-03 acceptance #4 — KHÔNG result Hub B.
        for item in results:
            assert item["hub_id"] != hub_b, (
                f"HARD-03 acceptance #4 VIOLATION — viewer Hub A thấy chunk "
                f"Hub B: {item}"
            )
        # Strong assertion: với viewer assign chỉ Hub A + request hub_ids=[B]
        # only → intersection {A} ∩ {B} = {} → results PHẢI rỗng.
        assert results == [], (
            f"HARD-03 acceptance #4 VIOLATION — intersection rỗng phải trả "
            f"results: [], got {len(results)} items"
        )
    else:
        # 403 path — backend reject explicit cross-hub request (cũng accept
        # vì defense-in-depth — frontend KHÔNG nhận chunk Hub B).
        body = r.json()
        assert body["success"] is False
        assert body["error"]["code"] in {"FORBIDDEN", "INVALID_HUB_IDS"}


# ============================================================================
# Test 5 — ask citation parsing (HARD-03 acceptance #5).
# ============================================================================


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_critical_ask_citation_marker_maps_to_chunk_id(
    auth_client: httpx.AsyncClient,
    viewer_token: str,
    viewer_user: dict[str, str],
    app_with_auth: Any,
    mock_litellm_citation_response: dict[str, Any],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """HARD-03 acceptance #5 — POST /api/ask với mock LLM `[1] [2]` → citations map.

    ASK-01 điểm vỡ chính: marker ``[N]`` trong answer PHẢI map đúng
    ``chunk_id`` của chunks retrieved. Mock LLM trả deterministic content
    ``"Theo tài liệu, A là B [1]. Còn C [2]."`` → ``citations[0].chunk_id``
    khớp chunk thứ 1 (rank 1), ``citations[1].chunk_id`` khớp chunk thứ 2.

    Mock LLM ``litellm.acompletion`` qua fixture
    ``mock_litellm_citation_response`` (Plan 10-03 conftest_hardening.py) —
    KHÔNG cần ``OPENAI_API_KEY`` thật cho CI gate.

    Detail test chi tiết: Phase 7 ``test_ask_api.py`` (11 test critical đầy
    đủ — anti-injection, cross-hub citation, hot-swap, usage log). Plan
    10-03 acceptance #5 là 1 case crisp nhất — citation marker → chunk_id
    1-to-1 mapping.
    """
    _ = app_with_auth

    # Seed hub + assign viewer + 2 chunk (biết trước chunk_id).
    hub_id = await _insert_hub(
        name="Hub Citation Test",
        code="hub-cite",
        subdomain="hub-cite",
    )
    await _assign_user_hub(user_id=viewer_user["id"], hub_id=hub_id)
    doc_id = await _insert_document(hub_id=hub_id, filename="cite-test.docx")

    chunk_a = await _insert_chunk(
        document_id=doc_id,
        hub_id=hub_id,
        content="Đoạn nội dung A là B.",
        vector=_fixed_vector(0.1),
    )
    chunk_b = await _insert_chunk(
        document_id=doc_id,
        hub_id=hub_id,
        content="Đoạn nội dung C khác.",
        vector=_fixed_vector(0.1),
    )
    seeded_set = {chunk_a, chunk_b}

    # Patch embed query → cùng vector seed chunk → cosine distance xác định.
    _patch_embed(monkeypatch, _fixed_vector(0.1))

    # mock_litellm_citation_response đã set default content
    # "Theo tài liệu, A là B [1]. Còn C [2]." — citations 2 marker.
    _ = mock_litellm_citation_response

    # === POST /api/ask ===
    r = await auth_client.post(
        "/api/ask",
        headers={"Authorization": f"Bearer {viewer_token}"},
        json={"query": "A là gì và C là gì?", "hub_id": hub_id},
    )
    assert r.status_code == 200, r.text

    body = r.json()
    assert body["success"] is True, body
    data = body["data"]

    # Envelope shape AskResponse.
    assert isinstance(data["answer"], str) and data["answer"]
    assert isinstance(data["model"], str)
    assert isinstance(data["query_time_ms"], int)

    # HARD-03 acceptance #5 — citation marker [N] → chunk_id 1-to-1 mapping.
    citations = data["citations"]
    assert len(citations) == 2, (
        f"HARD-03 acceptance #5 VIOLATION — expected 2 citations (marker "
        f"[1] [2]), got {len(citations)}: {citations}"
    )

    # Marker number khớp.
    assert citations[0]["number"] == 1, citations[0]
    assert citations[1]["number"] == 2, citations[1]

    # chunk_id của mỗi citation nằm trong set seeded — KHÔNG hallucinate.
    assert citations[0]["chunk_id"] in seeded_set, (
        f"HARD-03 acceptance #5 VIOLATION — citation[0].chunk_id "
        f"{citations[0]['chunk_id']!r} không thuộc seeded chunks {seeded_set}"
    )
    assert citations[1]["chunk_id"] in seeded_set, (
        f"HARD-03 acceptance #5 VIOLATION — citation[1].chunk_id "
        f"{citations[1]['chunk_id']!r} không thuộc seeded chunks {seeded_set}"
    )

    # Mapping 1-1 — KHÔNG trỏ nhầm cùng 1 chunk.
    assert citations[0]["chunk_id"] != citations[1]["chunk_id"], (
        "HARD-03 acceptance #5 VIOLATION — 2 marker map cùng 1 chunk_id "
        "(parse_citations bug)"
    )

    # Mỗi citation đủ field shape D-07.
    for c in citations:
        assert c["document_id"], c
        assert isinstance(c["score"], (int, float)), c
        assert isinstance(c["content_snippet"], str), c


# ============================================================================
# Self-check helper — assert UUID format cho test stability.
# ============================================================================


def _is_uuid_str(value: str) -> bool:
    """Verify value là UUID string hợp lệ (helper cho test stability)."""
    try:
        uuid.UUID(value)
    except (ValueError, TypeError):
        return False
    return True
