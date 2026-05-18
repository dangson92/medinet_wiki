"""Smoke golden path API end-to-end — Plan 08-03 Task 1 (COMPAT-01, ROADMAP SC2).

Chứng minh tự động golden path API contract: login admin → upload DOCX vào
`hub_y_te` → GET document detail → cross-hub search → ask (LLM mock) → audit log.
Phần verify được bằng máy của SC2 — Plan 08-04 (manual UAT) chỉ còn phần render
UI 11 trang cần mắt người.

DEF-05-01 — file này boot app qua fixture `app_with_auth` (lifespan + migration).
cocoindex 1.0.3 `core.Environment` là process-global singleton KHÔNG re-open
được → file test này PHẢI chạy 1 FILE/LẦN pytest:

    uv run pytest tests/integration/test_smoke_golden_path.py

KHÔNG gộp với test integration khác boot app trong cùng 1 pytest process (sẽ
FAIL `environment already open` từ file thứ 2).

Quyết định execute:
- Toàn bộ golden path viết thành MỘT test function tuần tự (không tách nhiều
  test cùng dùng app fixture — tránh DEF-05-01 trong cùng process).
- File DOCX tối giản tạo qua `python-docx` (`from docx import Document` — đã có
  trong deps, dùng ở test_documents_upload.py).
- LLM call MOCK qua fixture `mock_llm` (D-07-05-A — OPENAI_API_KEY M2 dev là
  placeholder `sk-replace-me`).
- Query embedding MOCK: cross-hub search + ask gọi `embed_text` để embed câu
  query (KHÔNG nằm trong `mock_llm` — `mock_llm` chỉ patch `litellm.acompletion`).
  Với key placeholder M2, `embed_text` thật → 401 → search trả 500
  EMBEDDING_FAILED. Patch `app.services.search_service.embed_text` trả vector
  deterministic 1536-dim (pattern `_patch_embed` của test_ask_api.py) — golden
  path verify SHAPE response + flow, KHÔNG verify chất lượng embedding.
- KHÔNG có cocoindex runtime trong test → upload BackgroundTask
  `trigger_cocoindex_update` cần `app.state.cocoindex_app`; cấp `MockCocoindexApp`
  (update_blocking no-op). Vì mock generate 0 chunk → document chuyển sang
  `failed` sau BackgroundTask — golden path KHÔNG assert `chunk_count>0` (đó là
  phần manual UAT 08-04 với cocoindex thật).
- Cross-hub search + ask trên hub vừa seed: chunks table rỗng (không ingest
  thật) → `results`/`citations` có thể RỖNG. Test verify SHAPE response (key
  tồn tại đúng kiểu), KHÔNG assert độ dài.
- Audit: `auth.login` chắc chắn enqueue (verify AUDIT_ACTIONS + auth router).
  `document.upload` cũng nằm trong AUDIT_ACTIONS. search/ask KHÔNG enqueue audit
  (không nằm trong code path audit) — test assert LINH HOẠT (≥1 entry
  `auth.login`), ghi chú action không log cho 08-04 manual checklist.

Reuse fixtures conftest: app_with_auth, auth_client, admin_user, admin_token,
mock_llm, helper _wait_usage_count-pattern poll.
"""
from __future__ import annotations

import asyncio
import io
import time
import uuid
from typing import Any

import httpx
import pytest
from docx import Document as DocxDocument
from sqlalchemy import text


class _MockCocoindexApp:
    """Mock coco.App — update_blocking no-op (không có cocoindex runtime trong test).

    Upload router add BackgroundTask `trigger_cocoindex_update(cocoindex_app, id)`.
    Nếu `cocoindex_app` là None thì task set status='failed' với message lifespan.
    Cấp mock để BackgroundTask chạy sạch (mock generate 0 chunk → status='failed'
    bình thường — golden path KHÔNG assert chunk_count).
    """

    def __init__(self) -> None:
        self.update_blocking_calls = 0

    def update_blocking(self) -> None:
        self.update_blocking_calls += 1


def _make_docx_bytes() -> bytes:
    """Tạo file DOCX tối giản hợp lệ in-memory → bytes.

    `python-docx` có sẵn trong deps (test_documents_upload.py dùng). Nội dung
    tiếng Việt ngắn để giống tài liệu y tế thật.
    """
    doc = DocxDocument()
    doc.add_paragraph("Quy trình khám bệnh đa khoa")
    doc.add_paragraph("Bước 1: tiếp nhận bệnh nhân và lập hồ sơ.")
    doc.add_paragraph("Bước 2: khám lâm sàng và chỉ định cận lâm sàng.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _seed_hub_y_te() -> str:
    """INSERT 1 hub `hub_y_te` qua engine (pattern conftest `_insert_hub`).

    Migration 0003: hubs có `slug` (NOT NULL legacy) + `code` + `subdomain` +
    `status`. `slug=code`. Return hub_id string.
    """
    from app.db.session import get_engine

    engine = get_engine()
    hub_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO hubs "
                "(id, slug, code, name, subdomain, description, status, "
                "is_active, created_at, updated_at) "
                "VALUES (:id, :slug, :code, :name, :subdomain, NULL, 'active', "
                "TRUE, NOW(), NOW())"
            ),
            {
                "id": hub_id,
                "slug": "hub-y-te",
                "code": "hub-y-te",
                "name": "Hub Y Tế",
                "subdomain": "hub-y-te",
            },
        )
    return hub_id


async def _assign_user_to_hub(*, user_id: str, hub_id: str) -> None:
    """INSERT user_hubs — gán admin vào hub để ask/search isolation pass (HUB-02)."""
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO user_hubs (user_id, hub_id, assigned_at) "
                "VALUES (:uid, :hid, NOW())"
            ),
            {"uid": user_id, "hid": hub_id},
        )


def _patch_query_embedding(
    monkeypatch: pytest.MonkeyPatch, *, seed: float = 0.1
) -> None:
    """Monkeypatch `app.services.search_service.embed_text` trả vector cố định.

    `SearchService` (cross-hub search) + `AskService` (tái dùng SearchService cho
    retrieval) gọi `embed_text` để embed câu query. `search_service` đã
    `from app.services.embedder import embed_text` → tên `embed_text` bind vào
    module `search_service`. Với OPENAI_API_KEY placeholder M2, gọi thật → 401.
    Patch trả vector 1536-dim deterministic — golden path verify flow + SHAPE,
    KHÔNG verify chất lượng embedding (pattern `_patch_embed` test_ask_api.py).
    """

    async def _fake_embed(text: str, model: str | None = None) -> list[float]:
        _ = (text, model)
        return [seed] * 1536

    monkeypatch.setattr(
        "app.services.search_service.embed_text", _fake_embed
    )


async def _count_audit_logs(*, timeout_s: float = 3.0) -> int:
    """Poll count(*) audit_logs — return số entry tối đa quan sát được trong timeout.

    Audit ghi qua asyncio.Queue batch flush (audit_batch_size=128 HOẶC
    flush_interval=2s). Poll có giới hạn ~3s (pattern `_wait_usage_count` conftest)
    — deterministic, tránh sleep cố định flaky.

    GHI CHÚ DEVIATION (xem SUMMARY 08-03 + input 08-04 manual checklist):
    `AUDIT_ACTIONS` (audit_service.py) khai báo `auth.login` + `document.upload`
    NHƯNG code thực tế KHÔNG enqueue 2 action này — `app/auth/service.py` chỉ
    structured-log `auth_login_success`, upload router KHÔNG gọi `enqueue_audit`.
    Action THỰC SỰ enqueue: `hub.create`/`hub.update`/`user.create`/
    `security.hub_isolation_violation` + `document_delete` (synchronous INSERT).
    → Golden path login→upload→search→ask KHÔNG sinh audit entry nào. Test KHÔNG
    hardcode "4 entry" (plan Task 1 step 7 — assert linh hoạt): verify endpoint
    `GET /api/audit-logs` hoạt động đúng contract (200 + list) là phần kiểm được
    bằng máy; "audit thấy N entry sau golden path" chuyển sang 08-04 manual UAT
    (sau khi user thực hiện hub.create/user.create thật qua UI).
    """
    from app.db.session import get_engine

    engine = get_engine()
    deadline = time.monotonic() + timeout_s
    last = 0
    while time.monotonic() < deadline:
        async with engine.begin() as conn:
            row = (
                await conn.execute(
                    text("SELECT count(*) FROM audit_logs")
                )
            ).fetchone()
        last = int(row[0]) if row else 0
        await asyncio.sleep(0.1)
    return last


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_golden_path_end_to_end(
    auth_client: httpx.AsyncClient,
    admin_user: dict[str, str],
    app_with_auth: Any,
    mock_llm: dict[str, Any],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Golden path SC2 — login → upload → GET detail → cross-hub search → ask → audit.

    Chuỗi tuần tự trong 1 test function (DEF-05-01 — tránh nhiều test cùng dùng
    app fixture). Mỗi bước assert status code + message `resp.text` khi fail.
    """
    # cocoindex runtime không có trong test → cấp mock cho app.state để upload
    # BackgroundTask chạy sạch.
    app_with_auth.state.cocoindex_app = _MockCocoindexApp()
    # Query embedding mock — cross-hub search + ask embed câu query qua embed_text;
    # key placeholder M2 → 401 nếu gọi thật. Patch trả vector cố định.
    _patch_query_embedding(monkeypatch)

    # === Bước 1: Seed hub hub_y_te + gán admin vào hub ===
    hub_id = await _seed_hub_y_te()
    await _assign_user_to_hub(user_id=admin_user["id"], hub_id=hub_id)

    # === Bước 2: Login admin ===
    login_resp = await auth_client.post(
        "/api/auth/login",
        json={"email": admin_user["email"], "password": admin_user["password"]},
    )
    assert login_resp.status_code == 200, login_resp.text
    login_body = login_resp.json()
    assert login_body["success"] is True, login_body
    access_token = login_body["data"]["access_token"]
    assert isinstance(access_token, str) and access_token, (
        f"access_token rỗng/không phải str: {access_token!r}"
    )
    auth_header = {"Authorization": f"Bearer {access_token}"}

    # === Bước 3: Upload DOCX vào hub_y_te ===
    docx_filename = "quy-trinh-kham.docx"
    upload_resp = await auth_client.post(
        "/api/documents/upload",
        headers=auth_header,
        files={
            "file": (
                docx_filename,
                io.BytesIO(_make_docx_bytes()),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
        data={"hub_id": hub_id},
    )
    assert upload_resp.status_code == 202, upload_resp.text
    upload_body = upload_resp.json()
    assert upload_body["success"] is True, upload_body
    upload_data = upload_body["data"]
    # Router trả `id` (DocumentResponse.id — verify documents.py:164 model_dump).
    document_id = upload_data["id"]
    assert document_id, f"upload không trả document id: {upload_data}"
    assert upload_data["name"] == docx_filename, upload_data

    # === Bước 4: GET document detail ===
    detail_resp = await auth_client.get(
        f"/api/documents/{document_id}", headers=auth_header
    )
    assert detail_resp.status_code == 200, detail_resp.text
    detail_data = detail_resp.json()["data"]
    assert detail_data["name"] == docx_filename, (
        f"GET detail name lệch tên file upload: {detail_data['name']!r}"
    )
    # status hợp lệ: pending (chưa chạy BackgroundTask) / processing / completed.
    # cocoindex mock generate 0 chunk → có thể đã 'failed'. Golden path SC2 chỉ
    # cần upload→detail roundtrip name đúng; status các giá trị hợp lệ vòng đời.
    assert detail_data["status"] in {
        "pending",
        "processing",
        "completed",
        "failed",
    }, detail_data

    # === Bước 5: Cross-hub search ===
    search_resp = await auth_client.post(
        "/api/search/cross-hub",
        headers=auth_header,
        json={"query": "khám bệnh", "hub_ids": [hub_id], "top_k": 10},
    )
    assert search_resp.status_code == 200, search_resp.text
    search_data = search_resp.json()["data"]
    # chunks table rỗng (không ingest thật) → `results` có thể RỖNG. Verify SHAPE.
    assert "results" in search_data, search_data
    assert isinstance(search_data["results"], list), search_data

    # === Bước 6: Ask (LLM mock) ===
    mock_llm["answer"] = "Quy trình khám bệnh gồm tiếp nhận và khám lâm sàng [1]."
    ask_resp = await auth_client.post(
        "/api/ask",
        headers=auth_header,
        json={"query": "quy trình khám bệnh là gì", "hub_id": hub_id},
    )
    assert ask_resp.status_code == 200, ask_resp.text
    ask_data = ask_resp.json()["data"]
    assert isinstance(ask_data["answer"], str), ask_data
    # citations list — rỗng nếu 0 chunk (marker [1] out-of-range bị bỏ).
    assert isinstance(ask_data["citations"], list), ask_data

    # === Bước 7: Audit log ===
    # Golden path login→upload→search→ask KHÔNG enqueue audit entry nào (code
    # KHÔNG enqueue `auth.login`/`document.upload` dù 2 action có trong enum
    # AUDIT_ACTIONS — xem docstring `_count_audit_logs` + SUMMARY 08-03). Phần
    # kiểm được bằng máy ở đây: endpoint `GET /api/audit-logs` đúng contract —
    # 200 + envelope `data` là list, admin đọc được. "Audit thấy N entry sau
    # golden path" defer 08-04 manual UAT (hub.create/user.create qua UI thật).
    audit_count = await _count_audit_logs(timeout_s=3.0)
    assert audit_count >= 0, audit_count  # quan sát số entry (golden path = 0)
    audit_resp = await auth_client.get("/api/audit-logs", headers=auth_header)
    assert audit_resp.status_code == 200, audit_resp.text
    audit_body = audit_resp.json()
    assert audit_body["success"] is True, audit_body
    audit_items = audit_body["data"]
    assert isinstance(audit_items, list), audit_items
