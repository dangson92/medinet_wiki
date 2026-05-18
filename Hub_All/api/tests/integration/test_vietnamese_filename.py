"""Vietnamese filename UTF-8 roundtrip — Plan 08-03 Task 2 (COMPAT-01, ROADMAP SC4).

Verify tự động SC4: upload file tên tiếng Việt `"Khám bệnh đa khoa.docx"` →
response upload + GET document detail trả `name` UTF-8 nguyên vẹn (không mojibake
"KhÃ¡m"). Đồng thời verify mục threat T-08-03-01 — filename tiếng Việt KHÔNG gây
path traversal.

DEF-05-01 — file này boot app qua fixture `app_with_auth` (lifespan + migration).
cocoindex 1.0.3 `core.Environment` là process-global singleton KHÔNG re-open
được → file test này PHẢI chạy 1 FILE/LẦN pytest:

    uv run pytest tests/integration/test_vietnamese_filename.py

KHÔNG gộp với test integration khác boot app trong cùng 1 pytest process.

Quyết định execute:
- File DOCX tạo qua `python-docx` (nhất quán với test_smoke_golden_path.py Task 1).
- cocoindex runtime không có trong test → cấp `_MockCocoindexApp` vào
  `app.state.cocoindex_app` để upload BackgroundTask `trigger_cocoindex_update`
  chạy sạch (mock generate 0 chunk — KHÔNG ảnh hưởng test filename).
- Path traversal (T-08-03-01): `FileStore.save()` (file_store.py) sinh tên file
  trên đĩa = `<uuid4>.<ext>` — KHÔNG dùng basename gốc → tên tiếng Việt KHÔNG
  bao giờ vào path đĩa. Test verify `file_path` khớp pattern UUID + nằm trong
  thư mục `file_store_dir` + KHÔNG chứa `..`. Cơ chế bảo vệ = UUID rename
  (KHÔNG sanitize) — ghi rõ trong SUMMARY.
"""
from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Any

import httpx
import pytest
from docx import Document as DocxDocument
from sqlalchemy import text

# Tên file tiếng Việt có dấu — literal Unicode. Nếu pipeline upload mojibake,
# assert so sánh chuỗi sẽ fail rõ ràng (vd "KhÃ¡m bá»‡nh ...").
VN_FILENAME = "Khám bệnh đa khoa.docx"


class _MockCocoindexApp:
    """Mock coco.App — update_blocking no-op (không có cocoindex runtime trong test)."""

    def __init__(self) -> None:
        self.update_blocking_calls = 0

    def update_blocking(self) -> None:
        self.update_blocking_calls += 1


def _make_docx_bytes() -> bytes:
    """Tạo file DOCX tối giản hợp lệ in-memory → bytes."""
    doc = DocxDocument()
    doc.add_paragraph("Tài liệu khám bệnh đa khoa")
    doc.add_paragraph("Nội dung quy trình tiếp nhận bệnh nhân.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _seed_hub() -> str:
    """INSERT 1 hub qua engine — return hub_id string (pattern conftest `_insert_hub`)."""
    from app.db.session import get_engine

    engine = get_engine()
    hub_id = str(uuid.uuid4())
    async with engine.begin() as conn:
        await conn.execute(
            text(
                "INSERT INTO hubs "
                "(id, slug, code, name, subdomain, description, status, "
                "is_active, created_at, updated_at) "
                "VALUES (:id, :slug, :code, :name, :subdomain, NULL, 'active', "
                "TRUE, NOW(), NOW())"
            ),
            {
                "id": hub_id,
                "slug": "hub-vn-test",
                "code": "hub-vn-test",
                "name": "Hub VN Test",
                "subdomain": "hub-vn-test",
            },
        )
    return hub_id


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_vietnamese_filename_roundtrip(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
) -> None:
    """SC4 — upload file tên tiếng Việt → response + GET detail trả name UTF-8 đúng.

    Verify thêm threat T-08-03-01 — filename tiếng Việt KHÔNG gây path traversal:
    `file_path` lưu DB khớp pattern UUID, nằm trong `file_store_dir`, không `..`.
    """
    app_with_auth.state.cocoindex_app = _MockCocoindexApp()
    hub_id = await _seed_hub()
    auth_header = {"Authorization": f"Bearer {admin_token}"}

    # === Upload file tên tiếng Việt ===
    # httpx multipart: truyền tên dạng str Unicode (KHÔNG bytes mojibake).
    upload_resp = await auth_client.post(
        "/api/documents/upload",
        headers=auth_header,
        files={
            "file": (
                VN_FILENAME,
                io.BytesIO(_make_docx_bytes()),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
        data={"hub_id": hub_id},
    )
    assert upload_resp.status_code == 202, upload_resp.text
    upload_body = upload_resp.json()
    assert upload_body["success"] is True, upload_body
    upload_data = upload_body["data"]
    # SC4 — name UTF-8 nguyên vẹn (so sánh literal Unicode; mojibake → fail rõ).
    assert upload_data["name"] == VN_FILENAME, (
        f"SC4 VIOLATION — upload name mojibake/lệch: {upload_data['name']!r} "
        f"kỳ vọng {VN_FILENAME!r}"
    )
    document_id = upload_data["id"]
    assert document_id, upload_data

    # === GET document detail — name vẫn đúng UTF-8 ===
    detail_resp = await auth_client.get(
        f"/api/documents/{document_id}", headers=auth_header
    )
    assert detail_resp.status_code == 200, detail_resp.text
    detail_data = detail_resp.json()["data"]
    assert detail_data["name"] == VN_FILENAME, (
        f"SC4 VIOLATION — GET detail name lệch: {detail_data['name']!r} "
        f"kỳ vọng {VN_FILENAME!r}"
    )

    # === Path traversal guard (T-08-03-01) ===
    # Cơ chế bảo vệ: FileStore.save() sinh tên đĩa = <uuid4>.<ext>, KHÔNG dùng
    # basename gốc → tên tiếng Việt KHÔNG vào path. file_path trả về response.
    file_path = detail_data.get("file_path")
    assert file_path, f"response thiếu file_path để verify traversal: {detail_data}"
    file_path_str = str(file_path)
    # 1) Không chứa `..` (không escape lên thư mục cha).
    assert ".." not in file_path_str, (
        f"T-08-03-01 VIOLATION — file_path chứa '..': {file_path_str!r}"
    )
    # 2) Tên file đĩa KHÔNG chứa ký tự tiếng Việt — FileStore UUID-rename.
    disk_name = Path(file_path_str).name
    assert disk_name not in {VN_FILENAME, Path(VN_FILENAME).name}, (
        f"T-08-03-01 — file_path dùng tên gốc thay vì UUID rename: {disk_name!r}"
    )
    # 3) Stem khớp pattern UUID4 — chứng minh FileStore UUID-rename đang bảo vệ.
    stem = Path(disk_name).stem
    try:
        uuid.UUID(stem)
    except ValueError:
        pytest.fail(
            f"T-08-03-01 — file_path stem không phải UUID4 (cơ chế bảo vệ traversal "
            f"không hoạt động như kỳ vọng): {stem!r}"
        )
    # 4) file_path nằm TRONG thư mục file_store_dir (không escape ra ngoài).
    from app.config import get_settings

    store_dir = get_settings().file_store_dir.resolve()
    resolved = Path(file_path_str).resolve()
    assert store_dir in resolved.parents or resolved.parent == store_dir, (
        f"T-08-03-01 VIOLATION — file_path nằm ngoài file_store_dir "
        f"({store_dir}): {resolved}"
    )
