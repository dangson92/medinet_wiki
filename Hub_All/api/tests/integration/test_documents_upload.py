"""Integration test Phase 4 Plan 04-04 REVISION 2 — POST /api/documents/upload + GET /:id.

Sử dụng fixtures Phase 3 Plan 03-05 từ tests/integration/conftest.py:
- postgres_container (scope=module), redis_container, app_with_auth, auth_client
- admin_user, viewer_user, admin_token, viewer_token

KHÔNG cần cocoindex flow chạy thật — mock app.state.cocoindex_app với
.update_blocking() sync method. Test focus vào router + service layer
(INGEST-04 + INGEST-05) + A4 BackgroundTasks pattern. Plan 04-06 sẽ test E2E
ingest pipeline với cocoindex thật.

REVISION 2 (A4 user BLOCKING confirmed):
- test 1 happy: assert mock cocoindex_app.update_blocking đã được call sau response.
- test 3 scanned PDF: assert mock cocoindex_app.update_blocking KHÔNG được call
  (BLOCKER #3 — scanned row final, KHÔNG add BackgroundTask).
- test 4 viewer 403: assert chỉ error.code (KHÔNG message text — WARNING #4).
- test 8: assert last_heartbeat IS NOT NULL sau upload (WARNING #7 bootstrap).
"""
from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Any

import httpx
import pytest
from docx import Document as DocxDocument
from sqlalchemy import text


class MockCocoindexApp:
    """Mock coco.App cho test — count update_blocking calls + thread-safe."""

    def __init__(self) -> None:
        self.update_blocking_calls: int = 0

    def update_blocking(self) -> None:
        # Sync method để asyncio.to_thread chạy được.
        self.update_blocking_calls += 1


def _make_docx_vn(tmp_path: Path, name: str = "Khám-bệnh.docx") -> bytes:
    """Tạo DOCX VN sample → bytes."""
    doc = DocxDocument()
    doc.add_paragraph("Mục 1. KHÁM TỔNG QUÁT")
    doc.add_paragraph("Bệnh nhân được khám lâm sàng tỉ mỉ.")
    path = tmp_path / name
    doc.save(str(path))
    return path.read_bytes()


async def _create_hub(app_with_auth: Any, hub_id: uuid.UUID | None = None) -> uuid.UUID:
    """INSERT 1 hub row trực tiếp (Phase 4 chưa có /api/hubs endpoint)."""
    _ = app_with_auth  # trigger lifespan + migration
    from app.db.session import get_engine

    hid = hub_id or uuid.uuid4()
    engine = get_engine()
    async with engine.begin() as conn:
        # Migration 0003 (Plan 05-01) thêm hubs.code / hubs.subdomain NOT NULL
        # (server_default đã drop) — helper phải truyền tường minh (DEF-05-02 fix).
        await conn.execute(
            text(
                "INSERT INTO hubs "
                "(id, slug, code, subdomain, name, description, "
                "is_active, created_at, updated_at) "
                "VALUES (:id, :slug, :code, :subdomain, :name, :desc, "
                "TRUE, NOW(), NOW())"
            ),
            {
                "id": str(hid),
                "slug": f"test-hub-{hid.hex[:8]}",
                "code": f"test-hub-{hid.hex[:8]}",
                "subdomain": f"test-hub-{hid.hex[:8]}",
                "name": "Test Hub",
                "desc": "Phase 4 test",
            },
        )
    return hid


@pytest.fixture
def mock_cocoindex_app(app_with_auth: Any) -> MockCocoindexApp:
    """Mock app.state.cocoindex_app cho A4 test — verify update_blocking gọi đúng."""
    mock = MockCocoindexApp()
    app_with_auth.state.cocoindex_app = mock
    return mock


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_happy_path_docx(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    admin_user: dict[str, str],
    app_with_auth: Any,
    tmp_path: Path,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """Admin upload DOCX VN → 202 + envelope shape D6. last_heartbeat bootstrap (WARNING #7).

    REVISION 2 A4: assert mock cocoindex_app.update_blocking đã được call qua
    BackgroundTask sau response. FastAPI BackgroundTasks chạy SAU response trả
    về client → cần await client.post() return rồi mới count.
    """
    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)

    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={
            "file": (
                "Khám-bệnh.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 202, r.text
    body = r.json()
    assert body["success"] is True
    assert body["error"] is None
    data = body["data"]
    assert "id" in data
    assert data["status"] == "pending"
    assert data["name"] == "Khám-bệnh.docx"

    # Verify row trong DB — last_heartbeat NOT NULL (WARNING #7 bootstrap).
    from app.db.session import get_engine
    engine = get_engine()
    async with engine.begin() as conn:
        row = (
            await conn.execute(
                text(
                    "SELECT status, hub_id, uploaded_by, filename, last_heartbeat "
                    "FROM documents WHERE id = :id"
                ),
                {"id": data["id"]},
            )
        ).fetchone()
    assert row is not None
    # Status có thể đã chuyển sang 'completed' hoặc 'failed' nếu BackgroundTask
    # đã chạy xong qua mock (mock generate 0 chunks → 'failed' với message).
    # Test core happy path chỉ verify row tồn tại + last_heartbeat bootstrap.
    assert row[0] in ("pending", "completed", "failed")
    assert str(row[1]) == str(hub_id)
    assert str(row[2]) == admin_user["id"]
    assert row[3] == "Khám-bệnh.docx"
    # WARNING #7: last_heartbeat bootstrap = NOW() — KHÔNG NULL.
    assert row[4] is not None, (
        "WARNING #7 violated — last_heartbeat phải bootstrap=NOW() lúc INSERT, KHÔNG NULL"
    )

    # A4 REVISION 2: verify BackgroundTask đã trigger update_blocking.
    # FastAPI BackgroundTasks chạy SAU response — chờ ngắn để task complete.
    import asyncio
    await asyncio.sleep(0.5)
    assert mock_cocoindex_app.update_blocking_calls >= 1, (
        f"A4 violated — update_blocking phải được call qua BackgroundTask, "
        f"got count={mock_cocoindex_app.update_blocking_calls}"
    )


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_unsupported_format(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """Upload .exe → 415 UNSUPPORTED_FORMAT (R4 mitigation). A4: KHÔNG trigger cocoindex."""
    hub_id = await _create_hub(app_with_auth)
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("malware.exe", io.BytesIO(b"MZ\x90\x00"), "application/octet-stream")},
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 415, r.text
    body = r.json()
    assert body["success"] is False
    assert body["error"]["code"] == "UNSUPPORTED_FORMAT"
    # WARNING #4: KHÔNG assert message text cụ thể (envelope có thể đổi).
    # Chỉ verify ext trong message để đảm bảo error contextual.
    assert ".exe" in body["error"]["message"]

    # A4: ext rejected SỚM (trước service.create) → KHÔNG add BackgroundTask.
    import asyncio
    await asyncio.sleep(0.3)
    assert mock_cocoindex_app.update_blocking_calls == 0, (
        f"A4 violated — .exe rejection KHÔNG được trigger cocoindex, "
        f"got count={mock_cocoindex_app.update_blocking_calls}"
    )


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_scanned_pdf(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    monkeypatch: pytest.MonkeyPatch,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """BLOCKER #3 — Upload scanned PDF → 415 UNSUPPORTED_FORMAT + status='failed_unsupported'.

    A4 REVISION 2: KHÔNG trigger cocoindex cho row scanned (final state).
    """
    hub_id = await _create_hub(app_with_auth)

    # Mock detect_scanned_pdf → return True (PDF scanned).
    # Patch module-level reference TRƯỚC khi service import-resolve.
    from app.services import documents_service, file_extract

    def _fake_detect(path: Path) -> bool:
        _ = path
        return True

    monkeypatch.setattr(file_extract, "detect_scanned_pdf", _fake_detect)
    monkeypatch.setattr(documents_service, "detect_scanned_pdf", _fake_detect)

    # Minimal PDF magic header (router save file, service detect mock → True).
    pdf_bytes = b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\nxref\n0 1\n0000000000 65535 f \ntrailer<<>>\n%%EOF"
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("scan.pdf", io.BytesIO(pdf_bytes), "application/pdf")},
        data={"hub_id": str(hub_id)},
    )
    # BLOCKER #3 — service detect scanned → INSERT failed_unsupported + raise → router 415.
    assert r.status_code == 415, r.text
    body = r.json()
    assert body["success"] is False
    assert body["error"]["code"] == "UNSUPPORTED_FORMAT"

    # Verify DB — phải có row status='failed_unsupported' (service INSERT trước khi raise).
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.connect() as conn:
        rows = (
            await conn.execute(
                text(
                    "SELECT status, filename, error_message FROM documents "
                    "WHERE hub_id = :hub AND filename = 'scan.pdf'"
                ),
                {"hub": str(hub_id)},
            )
        ).fetchall()
    assert len(rows) == 1, f"Phải có đúng 1 row scan.pdf trong DB, got {len(rows)}"
    assert rows[0][0] == "failed_unsupported", (
        f"BLOCKER #3 violated — status phải 'failed_unsupported', got {rows[0][0]}"
    )

    # A4 REVISION 2: scanned row final — router catch UnsupportedFormatError TRƯỚC
    # khi đến block add_task → KHÔNG trigger cocoindex.
    import asyncio
    await asyncio.sleep(0.3)
    assert mock_cocoindex_app.update_blocking_calls == 0, (
        f"A4 violated — scanned PDF KHÔNG được trigger cocoindex, "
        f"got count={mock_cocoindex_app.update_blocking_calls}"
    )


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_viewer_role(
    auth_client: httpx.AsyncClient,
    viewer_token: str,
    app_with_auth: Any,
    tmp_path: Path,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """Viewer KHÔNG được upload → 403 FORBIDDEN (RBAC require_role admin).

    WARNING #4: CHỈ assert error.code — KHÔNG assert message text cụ thể.
    """
    _ = mock_cocoindex_app  # fixture bind ensures app.state init even if test KHÔNG verify count.
    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {viewer_token}"},
        files={
            "file": (
                "ok.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 403, r.text
    body = r.json()
    assert body["success"] is False
    assert body["error"]["code"] == "FORBIDDEN"
    # WARNING #4: KHÔNG assert message text "Không đủ quyền..." — Phase 3 envelope format có thể đổi.


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_missing_auth(
    auth_client: httpx.AsyncClient,
    app_with_auth: Any,
    tmp_path: Path,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """No Bearer → 401 MISSING_AUTHORIZATION."""
    _ = mock_cocoindex_app  # fixture bind
    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)
    r = await auth_client.post(
        "/api/documents/upload",
        files={
            "file": (
                "ok.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 401, r.text
    body = r.json()
    assert body["error"]["code"] == "MISSING_AUTHORIZATION"


@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_invalid_hub_id(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    tmp_path: Path,
    app_with_auth: Any,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """hub_id KHÔNG phải UUID → 400 INVALID_HUB_ID."""
    _ = mock_cocoindex_app  # fixture bind
    _ = app_with_auth
    content = _make_docx_vn(tmp_path)
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={
            "file": (
                "x.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": "not-a-uuid"},
    )
    assert r.status_code == 400, r.text
    body = r.json()
    assert body["error"]["code"] == "INVALID_HUB_ID"


@pytest.mark.integration
@pytest.mark.asyncio
async def test_upload_rejects_empty_file(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """File rỗng → 400 EMPTY_FILE."""
    _ = mock_cocoindex_app  # fixture bind
    hub_id = await _create_hub(app_with_auth)
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("empty.txt", io.BytesIO(b""), "text/plain")},
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 400, r.text
    body = r.json()
    assert body["error"]["code"] == "EMPTY_FILE"


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_get_document_by_id_after_upload(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    tmp_path: Path,
    mock_cocoindex_app: MockCocoindexApp,
) -> None:
    """POST upload → 202; GET cùng id → 200 status='pending' + chunk_count=0 + last_heartbeat NOT NULL."""
    _ = mock_cocoindex_app  # fixture bind
    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={
            "file": (
                "test.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 202
    doc_id = r.json()["data"]["id"]

    r2 = await auth_client.get(
        f"/api/documents/{doc_id}",
        headers={"Authorization": f"Bearer {admin_token}"},
    )
    assert r2.status_code == 200, r2.text
    body = r2.json()
    assert body["success"] is True
    data = body["data"]
    assert data["id"] == doc_id
    # Status có thể đã chuyển sang 'failed' nếu BackgroundTask mock đã chạy
    # (mock generate 0 chunks → status='failed'). Accept cả 2 state.
    assert data["status"] in ("pending", "failed")
    assert data["name"] == "test.docx"
    # D6: payload KHÔNG còn expose attempts/last_heartbeat (Go shape không có).
    # WARNING #7 bootstrap last_heartbeat NOT NULL — verify trực tiếp DB.
    from app.db.session import get_engine
    engine = get_engine()
    async with engine.begin() as conn:
        row = (
            await conn.execute(
                text(
                    "SELECT attempts, last_heartbeat FROM documents WHERE id = :id"
                ),
                {"id": doc_id},
            )
        ).fetchone()
    assert row is not None
    assert row[0] == 0
    assert row[1] is not None, (
        "WARNING #7 violated — last_heartbeat phải bootstrap=NOW() lúc INSERT"
    )


@pytest.mark.integration
@pytest.mark.asyncio
async def test_get_document_404_unknown_id(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
) -> None:
    """GET unknown UUID → 404 NOT_FOUND."""
    _ = app_with_auth
    r = await auth_client.get(
        f"/api/documents/{uuid.uuid4()}",
        headers={"Authorization": f"Bearer {admin_token}"},
    )
    assert r.status_code == 404, r.text
    assert r.json()["error"]["code"] == "NOT_FOUND"
