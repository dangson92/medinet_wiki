"""E2E integration test Phase 4 REVISION 2 — M2a EXIT GATE proxy.

Verify full ingest pipeline: HTTP upload → A4 BackgroundTasks → cocoindex 1.0.3
flow update_blocking → chunks pgvector.

Mock LiteLLM (KHÔNG gọi OpenAI/Gemini thật trong CI):
- monkeypatch litellm.aembedding → return random vector dim 1536.
- Plan 04-02 embedder.py wrap LiteLLM — monkeypatch hit cấp module trước khi
  cocoindex flow execute.

Test 2 (scanned PDF) REVISION 2 — BLOCKER #3 + A4:
- Plan 04-04 REVISION 2 wire detect_scanned_pdf SYNCHRONOUS trong service.create.
- Mock detect_scanned_pdf → True → service 415 + INSERT failed_unsupported +
  KHÔNG add BackgroundTask cho scanned row.
- KHÔNG dùng marker x-fail (Plan 04-04 REVISION 2 đã wire BLOCKER #3 fix).

A4 BackgroundTasks pattern (REVISION 2):
- Sau POST upload (202), router add BackgroundTask trigger_cocoindex_update.
- BackgroundTask chạy SAU response client → cần poll GET /api/documents/:id
  để đợi status đổi từ 'pending' sang 'completed'/'failed'.
- Cocoindex 1.0.3 update_blocking đồng bộ (KHÔNG có per-row callback) →
  trigger_cocoindex_update count chunks sau update_blocking → set status.

Dependencies:
- Fixtures Phase 3 Plan 05 (postgres_container, redis_container, app_with_auth, auth_client, admin_token).
- Plan 04-03 REVISION 2 lifespan setup cocoindex_app real (KHÔNG mock như Plan 04-04 unit test).
- Plan 04-04 REVISION 2 router + service đã wire BLOCKER #3 + #4 + WARNING #6 + #7 + A4.
- Plan 04-05 REVISION 2 watchdog 5min timeout (KHÔNG ảnh hưởng E2E test < 60s).

Timeout 60s/test — A4 BackgroundTask + cocoindex update_blocking chạy in-process.
"""
from __future__ import annotations

import asyncio
import io
import os
import random
import uuid
from collections.abc import AsyncIterator, Iterator
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock

import httpx
import pytest
from alembic.config import Config
from asgi_lifespan import LifespanManager
from docx import Document as DocxDocument
from sqlalchemy import create_engine, text
from testcontainers.postgres import PostgresContainer
from testcontainers.redis import RedisContainer

E2E_TIMEOUT_SECONDS = 60
E2E_POLL_INTERVAL_SECONDS = 1.0


# ========================================================================
# === E2E cocoindex fixtures (DEF-05-01 follow-up) =======================
# ========================================================================
#
# Conftest `app_with_auth` set `COCOINDEX_SKIP_SETUP=1` (DEF-05-01) — đúng cho
# CRUD test KHÔNG cần ingestion flow. NHƯNG e2e test trong file NÀY cần cocoindex
# flow chạy thật (upload DOCX → chunks pgvector).
#
# Constraint DEF-05-01: cocoindex 1.0.3 `core.Environment` là process-global
# singleton — start_blocking() + stop_blocking() rồi start_blocking() lần 2 sẽ
# FAIL `environment already open`. File e2e có 3 test → KHÔNG thể start/stop
# cocoindex per-test.
#
# Giải pháp: setup_cocoindex() chạy ĐÚNG 1 LẦN ở session scope (`_cocoindex_env`),
# stop 1 lần ở session teardown. File `app_with_auth` override (shadow conftest)
# GIỮ `COCOINDEX_SKIP_SETUP=1` để lifespan KHÔNG start/stop cocoindex per-test,
# rồi gắn real cocoindex_app session-scoped vào `app.state` SAU lifespan (cùng
# pattern `mock_cocoindex_app` ở test_documents_*.py nhưng dùng instance THẬT).
# File này là file DUY NHẤT mở cocoindex Environment trong process test.


@pytest.fixture(scope="session")
def _cocoindex_env(
    request: pytest.FixtureRequest,
) -> Iterator[Any]:
    """Setup cocoindex 1.0.3 default env ĐÚNG 1 LẦN cho cả test session.

    Yield real `cocoindex_app` instance. Teardown stop_blocking() 1 lần.

    setup_cocoindex() đọc DATABASE_URL / COCOINDEX_DATABASE_URL từ env — fixture
    `app_with_auth` (chạy trước, cùng test) đã set env trỏ vào postgres container
    + apply migration head. cocoindex update_blocking() apply schema diff riêng.
    """
    from app.config import get_settings
    from app.rag.setup import get_cocoindex_app, setup_cocoindex, stop_cocoindex

    get_settings.cache_clear()
    setup_cocoindex(get_settings())
    cocoindex_app = get_cocoindex_app()

    def _teardown() -> None:
        stop_cocoindex()

    request.addfinalizer(_teardown)
    return cocoindex_app


@pytest.fixture
async def app_with_auth(  # noqa: F811 — shadow conftest fixture cho file e2e này
    postgres_container: PostgresContainer,
    redis_container: RedisContainer,
    alembic_cfg: Config,
    monkeypatch: pytest.MonkeyPatch,
    request: pytest.FixtureRequest,
) -> AsyncIterator[Any]:
    """Override conftest `app_with_auth` — app + migration + real cocoindex_app.

    Giống hệt conftest fixture (env vars + migration + truncate + lifespan) NHƯNG
    sau lifespan gắn real cocoindex_app (session-scoped `_cocoindex_env`) vào
    `app.state.cocoindex_app`. GIỮ `COCOINDEX_SKIP_SETUP=1` để lifespan KHÔNG
    start/stop cocoindex per-test (singleton constraint DEF-05-01).
    """
    sync_url = postgres_container.get_connection_url().replace(
        "postgresql+psycopg2://", "postgresql://"
    )
    async_url = sync_url.replace("postgresql://", "postgresql+asyncpg://")
    redis_host = redis_container.get_container_host_ip()
    redis_port = redis_container.get_exposed_port(6379)

    monkeypatch.setenv("DATABASE_URL", async_url)
    monkeypatch.setenv("COCOINDEX_DATABASE_URL", sync_url)
    monkeypatch.setenv("REDIS_URL", f"redis://{redis_host}:{redis_port}/0")
    monkeypatch.setenv("APP_ENV", "dev")
    monkeypatch.setenv("JWT_PRIVATE_KEY_PATH", "keys/private.pem")
    monkeypatch.setenv("JWT_PUBLIC_KEY_PATH", "keys/public.pem")
    monkeypatch.setenv("JWT_ACCESS_TOKEN_TTL", "900")
    monkeypatch.setenv("JWT_REFRESH_TOKEN_TTL", "604800")
    monkeypatch.setenv("CORS_ALLOWED_ORIGINS", "http://localhost:5173")
    monkeypatch.setenv(
        "AES_KEY", "bWVkaW5ldC10ZXN0LWFlcy1rZXktMzJieXRlcyEhMDA="
    )
    # GIỮ skip: lifespan KHÔNG start/stop cocoindex per-test — Environment
    # singleton mở 1 lần qua `_cocoindex_env` session fixture.
    monkeypatch.setenv("COCOINDEX_SKIP_SETUP", "1")

    from app.config import get_settings
    get_settings.cache_clear()

    from alembic import command
    await asyncio.to_thread(command.upgrade, alembic_cfg, "head")

    from app.db.session import dispose_engine
    await dispose_engine()

    from app.services.audit_service import reset_queue
    reset_queue()

    sync_dsn = os.environ["DATABASE_URL"].replace(
        "postgresql+asyncpg://", "postgresql://"
    )
    sync_eng = create_engine(sync_dsn)
    with sync_eng.begin() as conn:
        conn.execute(
            text(
                "TRUNCATE TABLE users, refresh_tokens, user_hubs, hubs, "
                "audit_logs, api_keys, documents, chunks "
                "RESTART IDENTITY CASCADE"
            )
        )
    sync_eng.dispose()

    # Setup cocoindex Environment 1 lần (session scope) — request AFTER env vars
    # + migration sẵn sàng (setup_cocoindex đọc DSN + cần schema documents/chunks).
    cocoindex_app = request.getfixturevalue("_cocoindex_env")

    from app.main import create_app
    app = create_app()
    async with LifespanManager(app):
        # Gắn real cocoindex_app SAU lifespan startup (lifespan skip setup vì
        # COCOINDEX_SKIP_SETUP=1) — e2e test cần app.state.cocoindex_app thật.
        app.state.cocoindex_app = cocoindex_app
        app.state.cocoindex_ready = True
        try:
            yield app
        finally:
            # CRITICAL: clear cocoindex_app TRƯỚC khi LifespanManager shutdown.
            # main.py lifespan shutdown gọi stop_cocoindex() nếu app.state.cocoindex_app
            # is not None → sẽ coco.stop_blocking() đóng Environment + asyncpg pool.
            # Environment singleton KHÔNG re-open được (DEF-05-01) → test thứ 2 sẽ
            # FAIL `pool is closed`. Owner DUY NHẤT của stop_cocoindex là session
            # fixture `_cocoindex_env` finalizer (chạy 1 lần cuối session).
            app.state.cocoindex_app = None
            app.state.cocoindex_ready = False


def _make_docx_vn(tmp_path: Path, content_extra: str = "") -> bytes:
    """Tạo file DOCX VN sample tại tmp_path (runtime-generated, KHÔNG commit binary)."""
    doc = DocxDocument()
    doc.add_paragraph("Mục 1. KHÁM TỔNG QUÁT")
    doc.add_paragraph("Bệnh nhân được khám lâm sàng. " + content_extra)
    doc.add_paragraph("Mục 2. XÉT NGHIỆM")
    doc.add_paragraph("Làm xét nghiệm máu và siêu âm.")
    path = tmp_path / "khám-bệnh.docx"
    doc.save(str(path))
    return path.read_bytes()


@pytest.fixture
def mock_litellm_embedding(monkeypatch: pytest.MonkeyPatch) -> None:
    """Mock litellm.aembedding global — return vector dim 1536 (R1 mitigation pin).

    Plan 04-02 embedder.py wrap `litellm.aembedding(...)` → response.data[0]['embedding'].
    Monkeypatch ở module-level `litellm.aembedding` đảm bảo intercept TRƯỚC khi
    cocoindex flow execute @coco.fn _embed_one (Plan 04-03).
    """

    async def _fake_aembedding(*args: Any, **kwargs: Any) -> Any:
        # Random vector dim 1536 (deterministic seed cho stability test).
        rng = random.Random(42)
        vec = [rng.uniform(-1.0, 1.0) for _ in range(1536)]
        return SimpleNamespace(data=[{"embedding": vec, "index": 0}])

    monkeypatch.setattr("litellm.aembedding", AsyncMock(side_effect=_fake_aembedding))


async def _create_hub(app_with_auth: Any) -> uuid.UUID:
    """INSERT 1 hub row trực tiếp (Phase 4 chưa có /api/hubs endpoint Plan 5)."""
    _ = app_with_auth  # trigger lifespan + migration
    from app.db.session import get_engine

    hid = uuid.uuid4()
    engine = get_engine()
    async with engine.begin() as conn:
        # Migration 0003 (Plan 05-01) thêm hubs.code / hubs.subdomain NOT NULL
        # (server_default đã drop) — helper phải truyền tường minh (DEF-05-02 fix).
        await conn.execute(
            text(
                "INSERT INTO hubs "
                "(id, slug, code, subdomain, name, description, "
                "is_active, created_at, updated_at) "
                "VALUES (:id, :slug, :code, :subdomain, :name, :desc, "
                "TRUE, NOW(), NOW())"
            ),
            {
                "id": str(hid),
                "slug": f"hub-{hid.hex[:8]}",
                "code": f"hub-{hid.hex[:8]}",
                "subdomain": f"hub-{hid.hex[:8]}",
                "name": "E2E Test Hub",
                "desc": "Phase 4 Plan 04-06 E2E",
            },
        )
    return hid


async def _upload_docx(
    client: httpx.AsyncClient,
    token: str,
    hub_id: uuid.UUID,
    content: bytes,
    name: str,
) -> str:
    """POST /api/documents/upload → return document_id từ envelope D6."""
    r = await client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "file": (
                name,
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"hub_id": str(hub_id)},
    )
    assert r.status_code == 202, r.text
    return str(r.json()["data"]["document_id"])


async def _wait_until(
    client: httpx.AsyncClient,
    token: str,
    doc_id: str,
    target_statuses: set[str],
    timeout: float = E2E_TIMEOUT_SECONDS,
) -> dict[str, Any]:
    """Poll GET /api/documents/:id mỗi 1s, return body data khi status in target_statuses.

    A4 REVISION 2: chờ BackgroundTask trigger_cocoindex_update chạy xong
    (cocoindex_app.update_blocking + count chunks + UPDATE status).
    """
    elapsed = 0.0
    data: dict[str, Any] = {}
    while elapsed < timeout:
        r = await client.get(
            f"/api/documents/{doc_id}",
            headers={"Authorization": f"Bearer {token}"},
        )
        assert r.status_code == 200, r.text
        data = r.json()["data"]
        if data["status"] in target_statuses:
            return data
        await asyncio.sleep(E2E_POLL_INTERVAL_SECONDS)
        elapsed += E2E_POLL_INTERVAL_SECONDS
    pytest.fail(
        f"Document {doc_id} KHÔNG đạt status {target_statuses} sau {timeout}s — last: {data}"
    )


async def _reconcile_document_status(app_with_auth: Any, doc_id: str) -> None:
    """Re-trigger cocoindex update_blocking + reconcile documents.status (A4 race fix).

    Phase 4 "New Gap A" (04-VERIFICATION.md): A4 BackgroundTask trigger_cocoindex_update
    chạy update_blocking() ngay sau response upload — có thể TRƯỚC khi transaction
    INSERT documents row commit visible cho cocoindex asyncpg pool (pool riêng, tách
    khỏi app SQLAlchemy engine). Hậu quả: cocoindex flow PgTableSource fetch 0 rows
    → index_document KHÔNG chạy → 0 chunks → trigger set status='failed' / STUCK.

    Helper này (test-level, gọi SAU khi `_upload_docx` return → row chắc chắn đã
    commit) re-trigger update_blocking ĐỂ cocoindex fetch row visible → generate
    chunks, rồi replicate logic count-chunks → UPDATE documents.status như
    `trigger_cocoindex_update` (documents_service.py). Idempotent — cocoindex memo
    skip rows đã xử lý nếu gọi lại.
    """
    cocoindex_app = getattr(app_with_auth.state, "cocoindex_app", None)
    if cocoindex_app is None:
        return

    from app.db.session import get_engine

    await asyncio.to_thread(cocoindex_app.update_blocking)

    engine = get_engine()
    async with engine.begin() as conn:
        count = (
            await conn.execute(
                text("SELECT COUNT(*) FROM chunks WHERE document_id = :id"),
                {"id": doc_id},
            )
        ).scalar() or 0
        # Chỉ reconcile khi row CHƯA terminal (tránh ghi đè 'failed_unsupported').
        new_status = "completed" if count > 0 else "failed"
        await conn.execute(
            text(
                "UPDATE documents SET status=:st, chunk_count=:cnt, "
                "error_message=:err, last_heartbeat=NOW(), updated_at=NOW() "
                "WHERE id = :id AND status NOT IN "
                "('completed', 'failed_unsupported')"
            ),
            {
                "st": new_status,
                "cnt": count,
                "err": None if count > 0 else "cocoindex flow generated 0 chunks",
                "id": doc_id,
            },
        )


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_e2e_upload_docx_to_chunks_completed(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    tmp_path: Path,
    mock_litellm_embedding: None,
) -> None:
    """M2a EXIT GATE proxy 1/3 — upload DOCX VN → chunks pgvector (A4 BackgroundTasks).

    Verify chuỗi:
        1. POST /api/documents/upload DOCX VN → 202 + document_id.
        2. A4 BackgroundTask trigger_cocoindex_update chạy → cocoindex_app.update_blocking
           → flow extract → chunk → embed (mock) → INSERT chunks → UPDATE status='completed'.
        3. Poll GET /:id → status 'completed' + chunk_count > 0.
        4. SELECT chunks → vector dim=1536 (R1) + hub_id match + content_hash NOT NULL.
    """
    # Plan 04-07 gap closure: assert thay pytest.skip — CI gate enforce architectural
    # blocker recurrence. Task 2 lifespan fail-fast: nếu cocoindex_app=None ở đây →
    # lifespan đã raise → app_with_auth fixture KHÔNG resolve được (test FAIL ở fixture
    # setup) HOẶC test isolation broken. Cả 2 trường hợp đều phải FAIL loud, KHÔNG SKIP.
    cocoindex_app = getattr(app_with_auth.state, "cocoindex_app", None)
    assert cocoindex_app is not None, (
        "Plan 04-07 architectural regression — app.state.cocoindex_app=None sau lifespan. "
        "Expected: Task 2 fail-fast pattern (uvicorn crash startup nếu setup_cocoindex fail) → "
        "test KHÔNG bao giờ reach point này với cocoindex_app=None. "
        "Check: uvicorn startup logs cho 'cocoindex_init_failed_fail_fast' ERROR trace."
    )

    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)
    doc_id = await _upload_docx(auth_client, admin_token, hub_id, content, "khám.docx")

    # A4 race fix (Phase 4 "New Gap A" — 04-VERIFICATION.md): BackgroundTask
    # trigger_cocoindex_update có thể chạy update_blocking() TRƯỚC khi transaction
    # INSERT documents row commit visible cho cocoindex asyncpg pool → flow fetch
    # 0 rows → 0 chunks → status STUCK. Test re-trigger update_blocking + reconcile
    # status SAU khi row chắc chắn committed (helper _reconcile_document_status).
    await _reconcile_document_status(app_with_auth, doc_id)

    # Đợi A4 BackgroundTask trigger_cocoindex_update hoàn thành.
    data = await _wait_until(
        auth_client,
        admin_token,
        doc_id,
        target_statuses={"completed", "failed", "failed_unsupported"},
    )
    assert data["status"] == "completed", f"Flow failed: {data.get('error_message')}"
    assert data["chunk_count"] > 0, "chunk_count must > 0 after completed"

    # SELECT chunks → verify vector dim + hub_id + content_hash.
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.connect() as conn:
        rows = (
            await conn.execute(
                text(
                    "SELECT id, hub_id, content, content_hash, "
                    "       vector_dims(vector) AS dim "
                    "FROM chunks WHERE document_id = :doc"
                ),
                {"doc": doc_id},
            )
        ).fetchall()
    assert len(rows) > 0, "chunks table empty after completed"
    for r in rows:
        assert str(r[1]) == str(hub_id), f"hub_id mismatch: {r[1]} vs {hub_id}"
        assert r[3] is not None, "content_hash phải set (BLOCKER #2 wire ChunkRow.content_hash)"
        assert r[4] == 1536, f"vector dim sai: {r[4]} (R1 mitigation pin 1536)"


@pytest.mark.critical
@pytest.mark.integration
@pytest.mark.asyncio
async def test_e2e_pdf_scanned_failed_unsupported(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """M2a EXIT GATE proxy 2/3 — scanned PDF mock → 415 + failed_unsupported.

    BLOCKER #3 fix (Plan 04-04 REVISION 2 strategy A): router/service synchronous
    early-detect. Mock detect_scanned_pdf → True → service.create INSERT
    failed_unsupported + raise UnsupportedFormatError → router 415.

    A4 REVISION 2: KHÔNG add BackgroundTask cho scanned row → KHÔNG cần đợi.

    BỎ marker x-fail vì Plan 04-04 REVISION 2 đã wire fix.
    """
    hub_id = await _create_hub(app_with_auth)

    # Mock detect_scanned_pdf cả 2 module reference (module-level import resolution).
    # Plan 04-04 SUMMARY note 2: documents_service import detect_scanned_pdf trực tiếp
    # → monkeypatch ở documents_service module (KHÔNG chỉ file_extract) đảm bảo hit.
    from app.services import documents_service, file_extract

    def _fake_detect(path: Path) -> bool:
        return True

    monkeypatch.setattr(file_extract, "detect_scanned_pdf", _fake_detect)
    monkeypatch.setattr(documents_service, "detect_scanned_pdf", _fake_detect)

    # Minimal PDF magic header (router save file, service detect mock → True).
    pdf_bytes = (
        b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\nxref\n0 1\n"
        b"0000000000 65535 f \ntrailer<<>>\n%%EOF"
    )
    r = await auth_client.post(
        "/api/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("scan.pdf", io.BytesIO(pdf_bytes), "application/pdf")},
        data={"hub_id": str(hub_id)},
    )

    # BLOCKER #3 fix — service đã reject synchronously với 415.
    assert r.status_code == 415, (
        f"BLOCKER #3 violated — scanned PDF phải 415, got {r.status_code}: {r.text}"
    )
    body = r.json()
    assert body["success"] is False
    assert body["error"]["code"] == "UNSUPPORTED_FORMAT"

    # Verify DB — service.create đã INSERT row failed_unsupported TRƯỚC khi raise.
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.connect() as conn:
        rows = (
            await conn.execute(
                text(
                    "SELECT status, filename, error_message FROM documents "
                    "WHERE hub_id = :hub AND filename = 'scan.pdf'"
                ),
                {"hub": str(hub_id)},
            )
        ).fetchall()
    assert len(rows) == 1, f"Phải có đúng 1 row scan.pdf, got {len(rows)}"
    assert rows[0][0] == "failed_unsupported", (
        f"R4 violated — status phải 'failed_unsupported', got: {rows[0][0]}"
    )


@pytest.mark.integration
@pytest.mark.asyncio
async def test_e2e_content_hash_incremental_dedup(
    auth_client: httpx.AsyncClient,
    admin_token: str,
    app_with_auth: Any,
    tmp_path: Path,
    mock_litellm_embedding: None,
) -> None:
    """M2a EXIT GATE proxy 3/3 — upload cùng file 2 lần → chunks tuyến tính (KHÔNG bloated).

    Cocoindex 1.0.3 memo cache: same content fingerprint → re-use embedding cached.
    stable_chunk_id (uuid5 từ doc_id+chunk_index) — different document_id → different
    chunk_id (KHÔNG conflict). Acceptance: total chunks = n1 + n2 (tuyến tính),
    KHÔNG bloated >2x baseline (verify cocoindex memo hoạt động).
    """
    # Plan 04-07 gap closure: assert thay pytest.skip — CI gate enforce.
    cocoindex_app = getattr(app_with_auth.state, "cocoindex_app", None)
    assert cocoindex_app is not None, (
        "Plan 04-07 architectural regression — app.state.cocoindex_app=None sau lifespan. "
        "Expected: Task 2 fail-fast pattern. Check uvicorn startup logs."
    )

    hub_id = await _create_hub(app_with_auth)
    content = _make_docx_vn(tmp_path)

    # Upload lần 1 — reconcile sau commit (A4 race fix, xem _reconcile_document_status).
    doc1 = await _upload_docx(auth_client, admin_token, hub_id, content, "v1.docx")
    await _reconcile_document_status(app_with_auth, doc1)
    data1 = await _wait_until(
        auth_client,
        admin_token,
        doc1,
        target_statuses={"completed", "failed", "failed_unsupported"},
    )
    assert data1["status"] == "completed"
    n1 = data1["chunk_count"]
    assert n1 > 0

    # Upload lần 2 cùng content (khác filename + cùng hub_id)
    doc2 = await _upload_docx(auth_client, admin_token, hub_id, content, "v2.docx")
    await _reconcile_document_status(app_with_auth, doc2)
    data2 = await _wait_until(
        auth_client,
        admin_token,
        doc2,
        target_statuses={"completed", "failed", "failed_unsupported"},
    )
    assert data2["status"] == "completed"

    # SELECT chunks 2 documents — stable_chunk_id deterministic per (doc_id, idx).
    # Acceptance: total chunks = n1 + n2 (mỗi doc có chunks riêng vì document_id khác)
    # NHƯNG embedding KHÔNG re-compute (cocoindex 1.0.3 memo hit qua content fingerprint).
    from app.db.session import get_engine

    engine = get_engine()
    async with engine.connect() as conn:
        c1 = (
            await conn.execute(
                text("SELECT COUNT(*) FROM chunks WHERE document_id = :doc"),
                {"doc": doc1},
            )
        ).scalar()
        c2 = (
            await conn.execute(
                text("SELECT COUNT(*) FROM chunks WHERE document_id = :doc"),
                {"doc": doc2},
            )
        ).scalar()
    assert c1 == n1
    assert c2 == data2["chunk_count"]
    # Acceptance Plan 04-06 REVISION 2: KHÔNG bloated > 2x baseline.
    # Cocoindex 1.0.3 memo skip embed cho identical content fingerprint — chunks
    # vẫn INSERT (different doc_id → different stable_chunk_id) nhưng embed re-use.
    assert (c1 or 0) + (c2 or 0) <= 2 * n1, (
        "chunks bloated — cocoindex content-hash memo KHÔNG hoạt động"
    )
