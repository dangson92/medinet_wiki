"""Sinh fixture PDF/DOCX deterministic cho pytest (I2 — fixtures commit vào git, không sinh runtime).

Chạy 1 lần:
    cd docling-pipeline
    python tests/fixtures/generate_fixtures.py

Sau đó commit 4 file PDF/DOCX vào git.
"""

from __future__ import annotations

from pathlib import Path

from PIL import Image as PILImage
from PIL import ImageDraw
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

FIX = Path(__file__).parent


def gen_sample_small_pdf() -> None:
    doc = SimpleDocTemplate(str(FIX / "sample_small.pdf"), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Hello World — sample small PDF for Docling tests.", styles["Title"]),
        Spacer(1, 12),
        Paragraph("This is a paragraph for extraction.", styles["Normal"]),
    ]
    doc.build(story)


def gen_sample_with_table_pdf() -> None:
    """B3 — PDF chứa 1 table 3x3 với header để test table_html preserved."""
    doc = SimpleDocTemplate(str(FIX / "sample_with_table.pdf"), pagesize=letter)
    styles = getSampleStyleSheet()
    data = [
        ["Cot A", "Cot B", "Cot C"],
        ["1", "2", "3"],
        ["4", "5", "6"],
    ]
    tbl = Table(data, colWidths=[100, 100, 100])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story = [Paragraph("Bang test", styles["Title"]), Spacer(1, 12), tbl]
    doc.build(story)


def gen_sample_with_figure_pdf() -> None:
    """B3 — PDF chứa 1 figure + caption để test figure caption marker."""
    img_path = FIX / "_tmp_figure.png"
    img = PILImage.new("RGB", (200, 100), color=(180, 220, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([20, 20, 180, 80], outline=(0, 0, 0), width=2)
    draw.text((40, 40), "FIGURE", fill=(0, 0, 0))
    img.save(img_path)

    doc = SimpleDocTemplate(str(FIX / "sample_with_figure.pdf"), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Tai lieu co hinh", styles["Title"]),
        Spacer(1, 12),
        Image(str(img_path), width=200, height=100),
        Paragraph("Hinh 1: So do minh hoa test", styles["Italic"]),
    ]
    doc.build(story)
    img_path.unlink()


def gen_sample_small_docx() -> None:
    from docx import Document

    d = Document()
    d.add_heading("Sample DOCX", 0)
    d.add_paragraph("Day la doan van test.")
    d.save(str(FIX / "sample_small.docx"))


if __name__ == "__main__":
    gen_sample_small_pdf()
    gen_sample_with_table_pdf()
    gen_sample_with_figure_pdf()
    gen_sample_small_docx()
    print("Generated fixtures in", FIX)
